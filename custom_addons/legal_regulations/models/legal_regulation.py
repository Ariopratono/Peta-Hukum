# -*- coding: utf-8 -*-
from odoo import models, fields, api
from datetime import date
import base64
import io
import re
import logging

_logger = logging.getLogger(__name__)


class LegalRegulation(models.Model):
    _name = 'legal.regulation'
    _description = 'Legal Regulation'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _order = 'hierarchy_order, tahun desc, nomor'
    _rec_name = 'judul'
    
    # Informasi Dasar - Diurutkan berdasarkan hierarki hukum Indonesia
    tipe_dokumen = fields.Selection([
        ('uud_1945', 'Undang-undang Dasar 1945'),
        ('tap_mpr', 'Ketetapan MPR'),
        ('undang_undang', 'Undang-Undang'),
        ('perpu', 'Peraturan Pemerintah Pengganti Undang-Undang'),
        ('peraturan_pemerintah', 'Peraturan Pemerintah'),
        ('peraturan_presiden', 'Peraturan Presiden'),
        ('keputusan_presiden', 'Keputusan Presiden'),
        ('instruksi_presiden', 'Instruksi Presiden'),
        ('peraturan_menteri', 'Peraturan Menteri'),
        ('keputusan_menteri', 'Keputusan Menteri'),
        ('peraturan_daerah', 'Peraturan Daerah'),
        ('peraturan_gubernur', 'Peraturan Gubernur'),
    ], string='Tipe Dokumen', required=True, default='undang_undang')
    
    # Field untuk menentukan urutan hierarki
    hierarchy_order = fields.Integer('Hierarchy Order', compute='_compute_hierarchy_order', store=True)
    
    judul = fields.Text('Judul', required=True, help='Judul lengkap peraturan')
    
    teu = fields.Char('T.E.U. (Tempat Terbit/Entitas Unit)', 
                      help='Contoh: Indonesia, Kementerian Pemberdayaan Perempuan dan Perlindungan Anak',
                      default='Indonesia')
    
    nomor = fields.Char('Nomor', required=True, help='Nomor peraturan')
    
    # Bentuk Peraturan
    bentuk = fields.Char('Bentuk', required=True, 
                         help='Contoh: Peraturan Menteri Pemberdayaan Perempuan Dan Perlindungan Anak')
    
    bentuk_singkat = fields.Char('Bentuk Singkat', 
                                 help='Contoh: Permen PPPA')
    
    # Tanggal dan Waktu
    tahun = fields.Integer('Tahun', required=True, default=lambda self: date.today().year)
    
    tempat_penetapan = fields.Char('Tempat Penetapan', default='Jakarta')
    
    tanggal_penetapan = fields.Date('Tanggal Penetapan', required=True)
    
    tanggal_pengundangan = fields.Date('Tanggal Pengundangan')
    
    tanggal_berlaku = fields.Date('Tanggal Berlaku')
    
    # Informasi Sumber dan Status
    sumber = fields.Text('Sumber', 
                         help='Contoh: BN 2025 (314); 70 hlm')
    
    subjek = fields.Text('Subjek', 
                         help='Contoh: KELUARGA, PERLINDUNGAN ANAK, PEREMPUAN / WANITA')
    
    status = fields.Selection([
        ('berlaku', 'Berlaku'),
        ('dicabut', 'Dicabut'),
        ('diubah', 'Diubah'),
        ('ditunda', 'Ditunda'),
        ('tidak_berlaku', 'Tidak Berlaku'),
    ], string='Status', required=True, default='berlaku')
    
    bahasa = fields.Selection([
        ('bahasa_indonesia', 'Bahasa Indonesia'),
        ('bahasa_inggris', 'Bahasa Inggris'),
        ('bahasa_daerah', 'Bahasa Daerah'),
    ], string='Bahasa', default='bahasa_indonesia')
    
    lokasi = fields.Char('Lokasi', 
                         help='Contoh: Kementerian Pemberdayaan Perempuan dan Perlindungan Anak')
    
    bidang = fields.Selection([
        ('hukum_pidana', 'HUKUM PIDANA'),
        ('hukum_perdata', 'HUKUM PERDATA'),
        ('hukum_administrasi_negara', 'HUKUM ADMINISTRASI NEGARA'),
        ('hukum_tata_negara', 'HUKUM TATA NEGARA'),
        ('hukum_internasional', 'HUKUM INTERNASIONAL'),
        ('hukum_bisnis', 'HUKUM BISNIS'),
        ('hukum_keluarga', 'HUKUM KELUARGA'),
        ('hukum_lingkungan', 'HUKUM LINGKUNGAN'),
        ('hukum_tenaga_kerja', 'HUKUM TENAGA KERJA'),
        ('hukum_pajak', 'HUKUM PAJAK'),
    ], string='Bidang Hukum')
    
    # Field Tambahan
    active = fields.Boolean('Active', default=True)
    
    keterangan = fields.Text('Keterangan Tambahan')
    
    # Field untuk Upload File PDF/TXT
    file_pdf = fields.Binary('File PDF', 
                             attachment=True,
                             help='File PDF peraturan (generated otomatis dari DOCX atau upload manual)')
    file_docx = fields.Binary('File DOCX (Word)', 
                              attachment=True,
                              help='Upload file DOCX peraturan (lebih akurat untuk extract text)')
    file_txt = fields.Binary('File TXT (Plain Text)',
                             attachment=True,
                             help='Upload file TXT (plain text) untuk ekstraksi teks yang lebih sederhana dan stabil')
    file_name = fields.Char('Nama File')
    file_size = fields.Integer('Ukuran File (KB)', compute='_compute_file_size', store=True)
    
    # Field untuk Full-Text Search
    isi_peraturan = fields.Html('Isi Peraturan', 
                               default='<p>Isi peraturan belum tersedia</p>',
                               help='Konten lengkap peraturan untuk pencarian mendalam')
    kata_kunci = fields.Text('Kata Kunci', 
                            default='',
                            help='Kata kunci tambahan untuk pencarian')
    ringkasan = fields.Text('Ringkasan', 
                           default='Ringkasan belum tersedia',
                           help='Ringkasan singkat peraturan')
    
    # Relasi Perubahan - untuk tracking UU yang mengubah UU ini
    perubahan_dari_ids = fields.One2many(
        'legal.regulation.perubahan',
        'peraturan_induk_id',
        string='Diubah Oleh',
        help='Daftar UU/Peraturan yang mengubah peraturan ini'
    )
    
    # Computed Fields
    nama_lengkap = fields.Char('Nama Lengkap', compute='_compute_nama_lengkap', store=True)
    
    is_berlaku_aktif = fields.Boolean('Berlaku Aktif', compute='_compute_is_berlaku_aktif')
    
    @api.depends('bentuk_singkat', 'nomor', 'tahun')
    def _compute_nama_lengkap(self):
        for record in self:
            if record.bentuk_singkat and record.nomor and record.tahun:
                record.nama_lengkap = f"{record.bentuk_singkat} No. {record.nomor} Tahun {record.tahun}"
            else:
                record.nama_lengkap = record.judul or 'Peraturan Baru'
    
    @api.depends('status', 'tanggal_berlaku')
    def _compute_is_berlaku_aktif(self):
        today = date.today()
        for record in self:
            if record.status == 'berlaku':
                if record.tanggal_berlaku:
                    record.is_berlaku_aktif = record.tanggal_berlaku <= today
                else:
                    record.is_berlaku_aktif = True
            else:
                record.is_berlaku_aktif = False
    
    @api.depends('file_pdf', 'file_docx', 'file_txt')
    def _compute_file_size(self):
        """Hitung ukuran file dalam KB"""
        for record in self:
            if record.file_txt:
                record.file_size = int(len(record.file_txt) / 1.37 / 1024)
            elif record.file_docx:
                # Priority DOCX
                record.file_size = int(len(record.file_docx) / 1.37 / 1024)
            elif record.file_pdf:
                record.file_size = int(len(record.file_pdf) / 1.37 / 1024)
            else:
                record.file_size = 0
    
    def _merge_broken_lines(self, text):
        """
        Merge lines that are incorrectly broken in the middle of sentences.
        Uses multi-pass approach to handle consecutive broken lines.
        IMPORTANT: Does NOT merge lines in Penjelasan section to preserve structure.
        Rules:
        1. Merge if line doesn't end with sentence terminators (., ;, :)
        2. Merge if next line continues numbering/legal references
        3. Don't merge if next line is a new section header (BAB, PASAL, etc.)
        4. Don't merge if next line starts with list markers ((1), a., etc.)
        5. Don't merge anything inside PENJELASAN section
        """
        import re
        
        _logger.info("=" * 80)
        _logger.info("MERGING BROKEN SENTENCES (MULTI-PASS)")
        _logger.info("=" * 80)
        
        # FIRST: Split text into main content and penjelasan section
        penjelasan_start_pattern = re.compile(r'^(PENJELASAN\s+ATAS|PENJELASAN|Penjelasan)\s*$', re.IGNORECASE)
        
        lines = text.split('\n')
        main_content_lines = []
        penjelasan_lines = []
        in_penjelasan = False
        
        for line in lines:
            if penjelasan_start_pattern.match(line.strip()):
                in_penjelasan = True
                penjelasan_lines.append(line)
            elif in_penjelasan:
                penjelasan_lines.append(line)
            else:
                main_content_lines.append(line)
        
        if penjelasan_lines:
            _logger.info(f"[OK] Found Penjelasan section with {len(penjelasan_lines)} lines - will NOT merge these")
            _logger.info(f"[OK] Main content has {len(main_content_lines)} lines - will merge these")
        
        # SECOND: Only merge main content (NOT penjelasan)
        text_to_merge = '\n'.join(main_content_lines)
        
        total_merges = 0
        pass_num = 0
        max_passes = 5  # Prevent infinite loop
        
        # Keep merging until no more merges possible (multi-pass)
        while pass_num < max_passes:
            pass_num += 1
            lines = text_to_merge.split('\n')
            merged_lines = []
            i = 0
            merge_count = 0
            
            _logger.info(f"Pass #{pass_num}: Processing {len(lines)} lines...")
            
            # Patterns that indicate a NEW section/item (should NOT be merged with previous line)
            new_section_patterns = [
                r'^\s*(BAB|BAGIAN|PARAGRAF)\s+',  # BAB I, BAGIAN PERTAMA, etc.
                r'^\s*Pasal\s+\d+[A-Z]?\s*$',  # Standalone "Pasal 1" or "Pasal 45A" (with optional letter suffix)
                r'^\s*Pasal\s+[IVX]+\s*$',  # Standalone "Pasal I", "Pasal II" (Roman numerals)
                r'^\s*\(\d+\)\s+\w',  # List items like (1) text
                r'^\s*[a-z]\.\s+\w',  # List items like a. text
                r'^\s*\d+\.\s+[A-Z]',  # List items like 1. Text (capitalized)
                r'^\s*(Mengingat|Menimbang|Memperhatikan|Menetapkan|Memutuskan|MEMUTUSKAN)\s*:',  # Legal headers
                r'^\s*DENGAN\s+RAHMAT',  # Opening phrase
                r'^\s*PRESIDEN\s+REPUBLIK',  # Title
                r'^\s*PENJELASAN',  # Penjelasan section
                r'^\s*I\.?\s+UMUM',  # I. UMUM
                r'^\s*II\.?\s+PASAL',  # II. PASAL DEMI PASAL
            ]
            
            # CRITICAL: Pattern for standalone Pasal that should NEVER be merged with next line
            # Support both numeric (1, 2, 45A) and Roman numerals (I, II, III, IV, V)
            standalone_pasal_pattern = re.compile(r'^\s*Pasal\s+(?:\d+[A-Z]?|[IVX]+)\s*$', re.IGNORECASE)
            
            # Sentence terminators - if line ends with these, it's complete
            sentence_terminators = re.compile(r'[.;:]\s*$')
        
            while i < len(lines):
                current_line = lines[i]
                current_stripped = current_line.strip()
                
                # Skip empty lines
                if not current_stripped:
                    merged_lines.append(current_line)
                    i += 1
                    continue
                
                # CRITICAL FIX: If current line is standalone Pasal (e.g., "Pasal 45B"), NEVER merge it
                if standalone_pasal_pattern.match(current_stripped):
                    merged_lines.append(current_line)
                    i += 1
                    _logger.info(f"  [OK] Preserving standalone Pasal header: '{current_stripped}'")
                    continue
                
                # Check if we should merge with next line
                should_merge = False
                
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    
                    # Skip if next line is empty
                    if not next_line:
                        merged_lines.append(current_line)
                        i += 1
                        continue
                    
                    # Check if next line is a new section (NEVER merge if true)
                    is_new_section = False
                    for pattern in new_section_patterns:
                        if re.match(pattern, next_line, re.IGNORECASE):
                            is_new_section = True
                            break
                    
                    if is_new_section:
                        # Next line is a new section - don't merge
                        merged_lines.append(current_line)
                        i += 1
                        continue
                    
                    # FIX: Jika line adalah nomor (ex: '8.') dan next line bukan header Pasal -> merge
                    if re.match(r'^\s*\d+\.\s*$', current_stripped):
                        if not re.match(r'^\s*Pasal\s+\d+[A-Z]?', next_line, re.IGNORECASE):
                            should_merge = True
                            merge_reason = "numeric index continuation"
                    
                    # Rule 6: Next line starts with number/letter that continues previous reference
                    # Example: "pasal 28F, pasal" (line break) "28G ayat (1)"
                    # CHECK THIS FIRST - highest priority for legal references
                    if re.match(r'^\s*\d+[A-Z]?\s+(ayat|huruf)', next_line, re.IGNORECASE):
                        should_merge = True
                        merge_reason = "reference continuation (ayat/huruf)"
                    
                    # Rule 7: Next line starts with "pasal" + number (continuing list of references)
                    # Example: "pasal 28F," (line break) "pasal 28G ayat (1)"
                    elif re.match(r'^\s*pasal\s+\d', next_line, re.IGNORECASE):
                        should_merge = True
                        merge_reason = "pasal reference continuation"
                    
                    # Rule 8: Next line starts with "dan pasal" or ", pasal"
                    elif re.match(r'^\s*(dan\s+)?pasal\s+\d', next_line, re.IGNORECASE):
                        should_merge = True
                        merge_reason = "dan pasal continuation"
                    
                    # Rule 5: Next line starts with continuation word (dan, atau, serta)
                    elif re.match(r'^\s*(dan|atau|serta|maupun)\s+', next_line, re.IGNORECASE):
                        should_merge = True
                        merge_reason = "continuation word"
                    
                    # Rule 2: Current line ends with comma and next continues
                    elif current_stripped.endswith(','):
                        should_merge = True
                        merge_reason = "comma continuation"
                    
                    # Rule 3: Special legal reference patterns
                    # Current line ends with incomplete reference marker
                    elif re.search(r'\b(pasal|ayat|huruf|nomor|tahun|undang-undang)\s*$', current_stripped, re.IGNORECASE):
                        should_merge = True
                        merge_reason = "incomplete legal reference"
                    
                    # Rule 4: Current line ends with open parenthesis or incomplete parenthesis
                    elif re.search(r'\([^)]*$', current_stripped):
                        should_merge = True
                        merge_reason = "incomplete parenthesis"
                    
                    # Rule 1: Current line doesn't end with sentence terminator (lowest priority)
                    elif not sentence_terminators.search(current_stripped):
                        should_merge = True
                        merge_reason = "no sentence terminator"
                    
                    # Perform merge if needed
                    if should_merge:
                        merged_line = current_line.rstrip() + ' ' + next_line
                        merged_lines.append(merged_line)
                        merge_count += 1
                        
                        # Log only first 5 merges per pass to avoid spam
                        if merge_count <= 5:
                            _logger.info(f"  [OK] Merge #{merge_count} ({merge_reason}) at line {i}:")
                            _logger.info(f"    '{current_stripped[:50]}...'")
                            _logger.info(f"    + '{next_line[:50]}...'")
                        
                        i += 2  # Skip next line since we merged it
                        continue
                
                # No merge - add current line as-is
                merged_lines.append(current_line)
                i += 1
            
            # Update text for next pass
            text_to_merge = '\n'.join(merged_lines)
            total_merges += merge_count
            
            _logger.info(f"  Pass #{pass_num} complete: {merge_count} merges")
            
            # If no merges in this pass, we're done
            if merge_count == 0:
                break
        
        _logger.info(f"[DONE] Sentence merging complete: {total_merges} total merges in {pass_num} passes")
        
        # THIRD: Reconstruct full text with merged main content + untouched penjelasan
        if penjelasan_lines:
            final_text = text_to_merge + '\n' + '\n'.join(penjelasan_lines)
            _logger.info(f"[OK] Reconstructed text: {len(text_to_merge.split(chr(10)))} main lines + {len(penjelasan_lines)} penjelasan lines")
        else:
            final_text = text_to_merge
            _logger.info(f"[OK] No Penjelasan section found, returning merged text only")
        
        _logger.info("=" * 80)
        
        return final_text

    def _fix_broken_words(self, text):
        """Fix words that are broken with spaces in the middle"""
        import re
        
        # Common Indonesian word patterns that might be broken
        
        # Pattern 1: Fix broken words with 1-3 letter chunks separated by single space
        def merge_short_chunks(text):
            # Find patterns like "xxx xx" or "xx xxx" where total length > 4
            pattern = r'\b([a-zA-Z]{2,4})\s([a-zA-Z]{2,5})\b'
            
            def check_and_merge(match):
                word1 = match.group(1)
                word2 = match.group(2)
                combined = word1 + word2
                
                # Common Indonesian prefixes that indicate word should be merged
                common_prefixes = ['per', 'ber', 'ter', 'men', 'mem', 'pen', 'pem', 'di', 'ke', 'se']
                # Common Indonesian suffixes
                common_suffixes = ['an', 'kan', 'nya', 'ku', 'mu', 'lah', 'kah', 'tah']
                
                # If word1 is a common prefix and combined makes sense
                if word1.lower() in common_prefixes:
                    return combined
                
                # If word2 is a common suffix
                if word2.lower() in common_suffixes:
                    return combined
                
                # If both are very short (2-3 chars) and combined length > 5, likely broken
                if len(word1) <= 3 and len(word2) <= 3 and len(combined) >= 5:
                    return combined
                
                # Otherwise keep original with space
                return match.group(0)
            
            return re.sub(pattern, check_and_merge, text)
        
        # Pattern 2: Fix broken words in ALL CAPS (existing fix)
        def fix_caps_words(match):
            word = match.group(0)
            if len(word.replace(' ', '')) > 5:
                return word.replace(' ', '')
            return word
        
        text = re.sub(r'\b([A-Z]\s){3,}[A-Z]\b', fix_caps_words, text)
        
        # Pattern 3: Fix specific common broken patterns
        # "xxx xx xxx" patterns (3+ segments)
        pattern_multi = r'\b([a-zA-Z]{2,4})\s([a-zA-Z]{1,3})\s([a-zA-Z]{2,5})\b'
        
        def merge_three_chunks(match):
            combined = match.group(1) + match.group(2) + match.group(3)
            # If middle segment is very short (1-2 chars), likely broken
            if len(match.group(2)) <= 2 and len(combined) >= 6:
                return combined
            return match.group(0)
        
        text = re.sub(pattern_multi, merge_three_chunks, text)
        
        # Apply the merge_short_chunks multiple times to catch nested patterns
        for _ in range(2):  # Run twice to catch consecutive broken words
            text = merge_short_chunks(text)
        
        return text
    
    def _is_standalone_pasal(self, stripped):
        """
        Return True jika baris adalah header Pasal VALID:
        Contoh valid:
            'Pasal 45'
            'PASAL 45A'
            'Pasal 5B'
            'Pasal I'
            'Pasal II'
            
        Return False jika:
            'PASAL 45B SETIAP ORANG ...'
            'Pasal 45 berbunyi sebagai berikut:'
            'Ketentuan Pasal 45A ...'
        """
        import re
        # cocokkan pasal yang berdiri sendiri (numeric atau Roman numeral)
        m = re.match(r'^(Pasal|PASAL)\s+(?:\d+[A-Z]?|[IVX]+)\s*$', stripped)
        return bool(m)
    
    def _format_text_to_html(self, text):
        """Format extracted text into properly structured HTML"""
        import re
        
        if not text or not text.strip():
            return ""
        
        # Fix broken words first
        text = self._fix_broken_words(text)
        
        lines = text.split('\n')
        html_parts = []
        current_paragraph = []
        in_list = False
        
        # Patterns untuk mendeteksi struktur
        # Pattern untuk numbering
        numbering_pattern = re.compile(r'^\s*(\d+\.|[a-z]\.|[A-Z]\.|[ivxIVX]+\.|[(\d]+[\).]|[(\w]+[\).])\s+(.+)$')
        # Pattern untuk BAB, Pasal (with optional letter suffix like 45A), Ayat, dll
        section_pattern = re.compile(r'^\s*(BAB|PASAL|Pasal\s+\d+[A-Z]?|Ayat|Bagian|Paragraf|BAGIAN|PARAGRAF)(\s+(.+))?$', re.IGNORECASE)
        # Pattern untuk header (ALL CAPS line)
        header_pattern = re.compile(r'^[A-Z\s]{10,}$')
        # Pattern untuk Mengingat/Menimbang - detect dengan atau tanpa titik dua
        # Akan di-check pertama untuk priority tinggi
        legal_intro_combined_pattern = re.compile(r'^\s*(Mengingat|Menimbang|Memperhatikan|Menetapkan|Memutuskan|Dengan Rahmat)\s*:?\s*(.*)$', re.IGNORECASE)
        
        # State for penjelasan block rendering
        in_penjelasan_block = False
        penjelasan_lines = []
        # Pattern for Pasal/Ayat/Huruf penjelasan
        penjelasan_header_pat = re.compile(r'^\s*💡 \[Penjelasan\s+(Pasal|Ayat|Huruf)[^\]]*\]:\s*$', re.IGNORECASE)
        # Pattern for Penjelasan Umum
        penjelasan_umum_pat = re.compile(r'^\s*📘 \[Penjelasan Umum\]:\s*$', re.IGNORECASE)
        # Pattern for alternative Penjelasan Umum format: "I. UMUM"
        penjelasan_umum_alternative_pat = re.compile(r'^\s*I\.\s+UMUM\s*$', re.IGNORECASE)
        # Pattern for subtitle "Dalam Undang-Undang ini yang dimaksud dengan:"
        subtitle_pattern = re.compile(r'^\s*Dalam\s+Undang-Undang\s+ini\s+yang\s+dimaksud\s+dengan:\s*$', re.IGNORECASE)

        def flush_current_paragraph():
            nonlocal current_paragraph, in_list
            if current_paragraph:
                if isinstance(current_paragraph[0], tuple):
                    para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                    indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                else:
                    para_text = ' '.join(current_paragraph)
                    indent_px = 0
                if in_list:
                    html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                else:
                    if indent_px > 0:
                        html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                    else:
                        html_parts.append(f'<p>{para_text}</p>')
                current_paragraph = []

        def flush_penjelasan_block():
            nonlocal in_penjelasan_block, penjelasan_lines
            if not in_penjelasan_block:
                return
            # First line is header, rest is content
            header = penjelasan_lines[0].strip() if penjelasan_lines else ''
            # Join content lines
            content_raw = ' '.join([ln.strip() for ln in penjelasan_lines[1:]]).strip()
            
            # Split by sentence breaks (. followed by capital letter) and wrap in <p> tags
            # This creates paragraph spacing without breaking the blue box
            paragraphs = re.split(r'\.\s+(?=[A-Z])', content_raw)
            content_html = ''
            for para in paragraphs:
                if para.strip():
                    # Add period back if it was removed by split
                    para_text = para.strip()
                    if not para_text.endswith('.'):
                        para_text += '.'
                    content_html += f'<p style="margin: 8px 0;">{para_text}</p>'
            
            # Render styled block
            html_parts.append(
                '<div class="penjelasan" style="margin-top: 0.5rem; margin-bottom: 0.75rem; padding: 8px 12px; '
                'background: #f4f8ff; border-left: 4px solid #0d6efd; border-radius: 2px;">'
                f'<div style="font-weight: 600; color: #0d6efd; margin-bottom: 8px;">{header}</div>'
                f'{content_html if content_html else ""}'
                '</div>'
            )
            penjelasan_lines = []
            in_penjelasan_block = False

        for line in lines:
            # PRESERVE LEADING SPACES for layout preservation
            # Count leading spaces before stripping
            leading_spaces = len(line) - len(line.lstrip())
            stripped = line.strip()
            
            # Handle penjelasan block rendering
            if in_penjelasan_block:
                # End of block on empty line
                if not stripped:
                    flush_penjelasan_block()
                    # continue to next line without treating this as paragraph break
                    continue
                else:
                    penjelasan_lines.append(line)
                    continue

            # Start of a penjelasan block (either Pasal/Ayat/Huruf or Umum)
            if penjelasan_header_pat.match(stripped) or penjelasan_umum_pat.match(stripped) or penjelasan_umum_alternative_pat.match(stripped):
                # Flush any open paragraph or list before starting
                flush_current_paragraph()
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                in_penjelasan_block = True
                # For "I. UMUM" format, convert to standard format
                if penjelasan_umum_alternative_pat.match(stripped):
                    penjelasan_lines = ['📘 [Penjelasan Umum]:']
                else:
                    penjelasan_lines = [line]
                continue

            # Preserve empty lines as paragraph breaks
            if not stripped:
                # If we were inside penjelasan (should have been handled above), flush
                if in_penjelasan_block:
                    flush_penjelasan_block()
                    continue
                if current_paragraph:
                    # Extract text and indentation from tuples
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        # For list continuation, don't add margin (controlled by <ul>)
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                elif in_list:
                    # Empty line while in list - close the list
                    html_parts.append('</ul>')
                    in_list = False
                continue
            
            # SPECIAL CASE: Check if line contains legal intro word but doesn't start with it
            # E.g., "PRESIDEN REPUBLIK INDONESIA, Menimbang" or "SK No 190185 A Mengingat"
            legal_intro_words = ['Mengingat', 'Menimbang', 'Memperhatikan', 'Menetapkan', 'Memutuskan', 'Dengan Rahmat']
            split_handled = False
            for intro_word in legal_intro_words:
                # Check if intro word appears but not at start (case-insensitive)
                # Pattern 1: With comma before intro word
                intro_pattern_with_comma = re.compile(r'^(.+?)\s*,\s*(' + intro_word + r')\b', re.IGNORECASE)
                # Pattern 2: Without comma - intro word at end of line after some text
                intro_pattern_no_comma = re.compile(r'^(.+?)\s+(' + intro_word + r')\s*:?\s*$', re.IGNORECASE)
                
                split_match = intro_pattern_with_comma.match(stripped) or intro_pattern_no_comma.match(stripped)
                if split_match:
                    # Found intro word after other text - split the line
                    before_part = split_match.group(1).strip()
                    intro_part = split_match.group(2).strip()
                    
                    # Make sure before_part is not empty and intro_part is actually the intro word
                    # Also check that before_part doesn't look like it's part of the intro phrase
                    if not before_part or len(before_part) < 3:
                        break
                    
                    # Process the "before" part first
                    if current_paragraph:
                        current_paragraph.append((before_part, indent_px))
                    else:
                        current_paragraph = [(before_part, indent_px)]
                    
                    # Flush current paragraph
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px_para = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px_para = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        if indent_px_para > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px_para}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                    
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    
                    # Now add the intro word as heading
                    html_parts.append(f'<h6 class="mt-3 mb-2"><strong>{intro_part}:</strong></h6>')
                    split_handled = True
                    break  # Break from for loop
            
            if split_handled:
                continue  # Continue to next line
            
            # Check for legal document intro patterns (Mengingat, Menimbang, etc.)
            # Works with or without colon - HIGH PRIORITY check
            legal_intro_match = legal_intro_combined_pattern.match(stripped)
            if legal_intro_match:
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                intro_word = legal_intro_match.group(1)
                rest = legal_intro_match.group(2).strip() if legal_intro_match.group(2) else ''
                
                # Check if original had colon - preserve it in heading
                has_colon = ':' in line and line.index(':') < line.index(intro_word) + len(intro_word) + 3
                if has_colon or rest:
                    # Had colon or has content after - add colon in heading
                    html_parts.append(f'<h6 class="mt-3 mb-2"><strong>{intro_word}:</strong></h6>')
                else:
                    # No colon, standalone - add without colon
                    html_parts.append(f'<h6 class="mt-3 mb-2"><strong>{intro_word}</strong></h6>')
                
                # Check if there's numbering right after
                if rest and numbering_pattern.match(rest):
                    number_match = numbering_pattern.match(rest)
                    if number_match:
                        # Start new list
                        html_parts.append('<ul style="list-style-type: none; padding-left: 0; margin-left: 1.5rem;">')
                        in_list = True
                        number = number_match.group(1)
                        content = number_match.group(2)
                        # Use table display for perfect alignment
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; display: table; width: 100%;"><span style="display: table-cell; width: 2rem; font-weight: bold; vertical-align: top;">{number}</span><span style="display: table-cell; word-wrap: break-word;">{content}</span></li>')
                elif rest:
                    html_parts.append(f'<p>{rest}</p>')
                
                continue
            
            # Check for subtitle pattern "Dalam Undang-Undang ini yang dimaksud dengan:"
            subtitle_match = subtitle_pattern.match(stripped)
            if subtitle_match:
                if current_paragraph:
                    flush_current_paragraph()
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                # Render as italic paragraph (not bold, like Penjelasan Umum style)
                html_parts.append(f'<p style="font-style: italic; margin-top: 10px; margin-bottom: 10px; color: #555;">{stripped}</p>')
                continue
            
            # Check for section headers (BAB, PASAL, etc.)
            section_match = section_pattern.match(stripped)
            
            # --- FIX: Hanya treat sebagai header jika PASAL berdiri sendiri ---
            skip_section_header = False
            # Support both numeric (Pasal 1, Pasal 45A) and Roman (Pasal I, Pasal II)
            pasal_standalone_pattern = re.compile(r'^(Pasal|PASAL)\s+(?:\d+[A-Z]?|[IVX]+)\s*$', re.IGNORECASE)
            is_pasal_line = pasal_standalone_pattern.match(stripped)
            
            if section_match and is_pasal_line:
                # Ini adalah "Pasal 45" atau "Pasal 45B" standalone - treat as header
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                html_parts.append(f'<h5 class="mt-3 mb-2"><strong>{stripped}</strong></h5>')
                continue
            elif section_match and not is_pasal_line:
                # Section header lain (BAB, BAGIAN, dll) atau Pasal dengan teks tambahan
                # Jika Pasal dengan teks (ex: "Pasal 45B Setiap Orang"), treat sebagai paragraf
                if re.match(r'^(Pasal|PASAL)\s+\d+[A-Z]?\s+\w', stripped, re.IGNORECASE):
                    # "Pasal 45B Setiap..." - treat as paragraph, not header
                    current_paragraph.append((stripped, 0))
                    continue
                else:
                    # BAB, BAGIAN, etc - treat as header
                    if current_paragraph:
                        # Extract text and indentation
                        if isinstance(current_paragraph[0], tuple):
                            para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                            indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                        else:
                            para_text = ' '.join(current_paragraph)
                            indent_px = 0
                        
                        if in_list:
                            html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                        else:
                            # Only add margin if there's actual indentation
                            if indent_px > 0:
                                html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                            else:
                                html_parts.append(f'<p>{para_text}</p>')
                        current_paragraph = []
                    
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    
                    html_parts.append(f'<h5 class="mt-3 mb-2"><strong>{stripped}</strong></h5>')
                    continue
            
            # Check for ALL CAPS header
            if len(stripped) > 10 and header_pattern.match(stripped):
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                html_parts.append(f'<h6 class="mt-2 mb-2"><strong>{stripped}</strong></h6>')
                continue
            
            # Check for numbered/lettered lists - PRESERVE ORIGINAL FORMAT
            numbering_match = numbering_pattern.match(stripped)
            # Exception: "I. UMUM" should NOT be treated as numbering (will be handled by penjelasan check above)
            if numbering_match and not penjelasan_umum_alternative_pat.match(stripped):
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        para_indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        para_indent_px = 0
                    
                    if in_list:
                        # Apply hanging indent to multi-line content
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        html_parts.append(f'<p style="margin-left: {para_indent_px}px;">{para_text}</p>')
                    current_paragraph = []
                
                if not in_list:
                    # Start new list with fixed left margin (no nested indentation)
                    html_parts.append('<ul style="list-style-type: none; padding-left: 0; margin-left: 1.5rem;">')
                    in_list = True
                
                # Add list item with ORIGINAL numbering and hanging indent
                # Use table display for perfect alignment
                # For list items, DON'T add margin-left (already controlled by <ul>)
                # Just use table display for proper alignment
                number = numbering_match.group(1)
                content = numbering_match.group(2)
                html_parts.append(f'<li style="margin-bottom: 0.5rem; display: table; width: 100%;"><span style="display: table-cell; width: 2rem; font-weight: bold; vertical-align: top;">{number}</span><span style="display: table-cell; word-wrap: break-word;">{content}</span></li>')
                continue
            
            # If not a special pattern, accumulate into current paragraph
            # This preserves line breaks within sections
            # Store with indentation info
            if leading_spaces > 0:
                indent_px = leading_spaces * 8
                current_paragraph.append((stripped, indent_px))
            else:
                current_paragraph.append((stripped, 0))
        
        # Close any remaining content
        # Flush penjelasan block if still open
        if in_penjelasan_block:
            flush_penjelasan_block()
        if current_paragraph:
            # Extract text and indentation
            if isinstance(current_paragraph[0], tuple):
                para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
            else:
                para_text = ' '.join(current_paragraph)
                indent_px = 0
            
            if in_list:
                # Apply hanging indent to closing content
                html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                html_parts.append('</ul>')
            else:
                # Only add margin if there's actual indentation
                if indent_px > 0:
                    html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                else:
                    html_parts.append(f'<p>{para_text}</p>')
        elif in_list:
            html_parts.append('</ul>')
        
        return ''.join(html_parts)
    
    def _extract_text_from_pdf(self, pdf_data):
        """Extract text content from PDF file - extract plain text then format with _format_text_to_html"""
        try:
            # Decode base64 PDF data
            pdf_bytes = base64.b64decode(pdf_data)
            
            # Use PyMuPDF to extract plain text (consistent with TXT extraction)
            _logger.info("Extracting plain text from PDF via PyMuPDF...")
            import fitz
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text = "\n".join([page.get_text("text") for page in doc])
            doc.close()
            
            if text.strip():
                _logger.info(f"PyMuPDF extracted {len(text)} characters of plain text")
                
                # Format using the same _format_text_to_html used by TXT extraction
                formatted_text = self._format_text_to_html(text)
                
                result = f'<div class="txt-content" style="line-height: 1.8;">{formatted_text}</div>'
                return result
            else:
                _logger.warning("PyMuPDF returned empty text, trying PDFExtractor fallback...")
                # Fallback to PDFExtractor
                from .parser import PDFExtractor
                extractor = PDFExtractor(pdf_bytes)
                text = extractor.extract_text_with_layout()
                if not text.strip():
                    text = extractor.extract_text_auto()
                if text.strip():
                    formatted_text = self._format_text_to_html(text)
                    return f'<div class="txt-content" style="line-height: 1.8;">{formatted_text}</div>'
                return '<p class="text-warning"><strong>Perhatian:</strong> Tidak dapat mengekstrak teks dari PDF ini.</p>'
                
        except ImportError as ie:
            # Fallback to old method if parser.py or its dependencies not available
            _logger.warning(f"PDFExtractor not available ({ie}), falling back to basic extraction...")
            
            try:
                pdf_bytes = base64.b64decode(pdf_data)
                pdf_file = io.BytesIO(pdf_bytes)
                
                # Try PyPDF2 first
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(pdf_file)
                    html_content = ['<div class="pdf-content" style="line-height: 1.8;">']
                    
                    for page_num, page in enumerate(reader.pages):
                        try:
                            text = page.extract_text()
                            if text.strip():
                                html_content.append(f'<div class="page-break mb-4">')
                                html_content.append(f'<h4 class="text-primary border-bottom pb-2 mb-3">Halaman {page_num + 1}</h4>')
                                
                                # Format text with proper structure
                                formatted_text = self._format_text_to_html(text)
                                html_content.append(formatted_text)
                                html_content.append('</div>')
                        except Exception as e:
                            _logger.warning(f"Error extracting page {page_num + 1}: {e}")
                            continue
                    
                    html_content.append('</div>')
                    
                    if len(html_content) > 2:  # More than just opening and closing div
                        return ''.join(html_content)
                        
                except ImportError:
                    _logger.warning("PyPDF2 not installed, trying pdfminer.six...")
                    
                # Fallback to pdfminer.six
                try:
                    from pdfminer.high_level import extract_text
                    pdf_file.seek(0)  # Reset file pointer
                    text = extract_text(pdf_file)
                    
                    if text.strip():
                        # Format text with proper structure
                        formatted_text = self._format_text_to_html(text)
                        return f'<div class="pdf-content" style="line-height: 1.8;">{formatted_text}</div>'
                        
                except ImportError:
                    _logger.error("Neither PyPDF2 nor pdfminer.six is installed")
                    return '<p class="text-danger"><strong>Error:</strong> Library untuk ekstraksi PDF tidak tersedia. Silakan install PyPDF2 atau pdfminer.six</p>'
                
                return '<p class="text-warning"><strong>Perhatian:</strong> Tidak dapat mengekstrak teks dari PDF ini.</p>'
                
            except Exception as fallback_error:
                _logger.error(f"Fallback extraction failed: {fallback_error}")
                return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari PDF: {fallback_error}</p>'
            
        except Exception as e:
            _logger.error(f"Error extracting PDF text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari PDF: {str(e)}</p>'
    
    def _get_paragraph_xml_text(self, paragraph):
        """Get raw XML text from paragraph including numbering"""
        try:
            from lxml import etree
            
            # Get the paragraph's XML element
            p_element = paragraph._element
            
            # Try to get any numbering or list information from XML
            xml_str = etree.tostring(p_element, encoding='unicode', method='text')
            
            return xml_str.strip() if xml_str else None
            
        except Exception as e:
            _logger.debug(f"Could not get XML text: {str(e)}")
            return None
    
    def _detect_list_format(self, text):
        """Detect if text starts with list numbering and extract it"""
        import re
        
        if not text:
            return None, text
        
        # Patterns untuk berbagai format numbering
        patterns = [
            r'^([a-z]\.)\s+(.+)$',  # a. text
            r'^([a-z]\))\s+(.+)$',  # a) text
            r'^(\d+\.)\s+(.+)$',    # 1. text
            r'^(\d+\))\s+(.+)$',    # 1) text
            r'^([ivxlc]+\.)\s+(.+)$',  # i. text (roman)
            r'^(\([a-z]\))\s+(.+)$',  # (a) text
            r'^(\(\d+\))\s+(.+)$',   # (1) text
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text.strip(), re.IGNORECASE)
            if match:
                return match.group(1), match.group(2)
        
        return None, text
    
    def _extract_text_from_docx(self, docx_data):
        """Extract text from DOCX file - MUCH MORE ACCURATE than PDF!"""
        try:
            from docx import Document
            from lxml import etree
            import re
            
            # Decode base64 DOCX data
            docx_bytes = base64.b64decode(docx_data)
            docx_file = io.BytesIO(docx_bytes)
            
            doc = Document(docx_file)
            
            # Extract all paragraphs with proper structure
            all_text = []
            list_item_counter = {}  # Track list items per numbering ID
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    # Empty paragraph = line break
                    all_text.append('')
                    continue
                
                # Check if this paragraph has Word numbering format
                has_word_numbering = False
                generated_numbering = None
                
                try:
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.numPr
                        if numPr is not None:
                            # This paragraph has numbering format in Word
                            has_word_numbering = True
                            
                            # Try to get numbering level
                            ilvl = 0
                            numId = 0
                            if numPr.ilvl is not None:
                                ilvl = int(numPr.ilvl.val)
                            if numPr.numId is not None:
                                numId = int(numPr.numId.val)
                            
                            # Track counter for this numbering ID and level
                            counter_key = f"{numId}_{ilvl}"
                            if counter_key not in list_item_counter:
                                list_item_counter[counter_key] = 0
                            list_item_counter[counter_key] += 1
                            count = list_item_counter[counter_key]
                            
                            # Generate numbering based on level
                            if ilvl == 0:
                                # Lowercase letters for level 0
                                if count <= 26:
                                    generated_numbering = f"{chr(96 + count)}."
                                else:
                                    generated_numbering = f"{count}."
                            elif ilvl == 1:
                                # Numbers for level 1
                                generated_numbering = f"{count}."
                            else:
                                # Default to numbers
                                generated_numbering = f"{count}."
                            
                except Exception as e:
                    _logger.debug(f"Numbering detection from structure failed: {str(e)}")
                
                # Now decide what to do based on what we found
                if has_word_numbering and generated_numbering:
                    # This paragraph has Word numbering
                    # Check if the text ALREADY starts with the same or similar numbering
                    detected_num, remaining_text = self._detect_list_format(text)
                    
                    if detected_num:
                        # Text already has numbering - DON'T add generated numbering
                        all_text.append(text)
                        _logger.info(f"Paragraph has Word numbering BUT text already has it: {text[:50]}...")
                    else:
                        # Text doesn't have numbering yet - ADD generated numbering
                        full_text = f"{generated_numbering} {text}"
                        all_text.append(full_text)
                        _logger.info(f"Added generated numbering {generated_numbering} to: {text[:30]}...")
                else:
                    # No Word numbering detected
                    # Just use the text as-is (may or may not have manual numbering)
                    all_text.append(text)
            
            # IMPORTANT: Keep line breaks when joining
            # This preserves paragraph structure and list items
            full_text = '\n'.join(all_text)
            
            _logger.info(f"DOCX extraction: {len(all_text)} paragraphs extracted")
            _logger.info(f"First 500 chars: {full_text[:500]}")
            
            # Format using our smart formatter (will detect Menimbang:, numbering, etc.)
            formatted_text = self._format_text_to_html(full_text)
            
            result = f'<div class="docx-content" style="line-height: 1.8;">{formatted_text}</div>'
            
            if len(result) > 50:  # Has content
                _logger.info(f"DOCX extraction successful: {len(all_text)} paragraphs, {len(full_text)} chars")
                return result
            else:
                return '<p class="text-warning"><strong>Perhatian:</strong> File DOCX kosong atau tidak dapat dibaca.</p>'
                
        except ImportError as ie:
            missing_lib = str(ie).split("'")[1] if "'" in str(ie) else "unknown"
            _logger.error(f"Missing library: {missing_lib}")
            return f'<p class="text-danger"><strong>Error:</strong> Library {missing_lib} tidak terinstall.</p>'
        except Exception as e:
            _logger.error(f"Error extracting DOCX text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari DOCX: {str(e)}</p>'
    
    def _convert_docx_to_pdf(self, docx_data):
        """Convert DOCX to PDF using python-docx and reportlab"""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Decode DOCX
            docx_bytes = base64.b64decode(docx_data)
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                                       rightMargin=72, leftMargin=72,
                                       topMargin=72, bottomMargin=18)
            
            # Container for PDF elements
            elements = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            style_normal = ParagraphStyle('CustomNormal', parent=styles['Normal'],
                                         fontSize=11, leading=14, alignment=TA_JUSTIFY)
            style_heading = ParagraphStyle('CustomHeading', parent=styles['Heading1'],
                                          fontSize=14, leading=16, spaceAfter=12)
            
            # Extract and add content
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    elements.append(Spacer(1, 0.2*inch))
                    continue
                
                # Escape HTML special chars
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Check if heading (ALL CAPS or starts with BAB, PASAL, etc)
                if text.isupper() and len(text) > 10:
                    elements.append(Paragraph(f"<b>{text}</b>", style_heading))
                elif text.upper().startswith(('BAB ', 'PASAL ', 'AYAT ', 'BAGIAN ')):
                    elements.append(Paragraph(f"<b>{text}</b>", style_heading))
                else:
                    elements.append(Paragraph(text, style_normal))
            
            # Build PDF
            pdf_doc.build(elements)
            
            # Get PDF bytes and encode to base64
            pdf_bytes = pdf_buffer.getvalue()
            pdf_base64 = base64.b64encode(pdf_bytes)
            
            return pdf_base64
            
        except ImportError as e:
            _logger.error(f"Missing library for PDF conversion: {e}. Install: pip install python-docx reportlab")
            return None
        except Exception as e:
            _logger.error(f"Error converting DOCX to PDF: {str(e)}")
            return None
    
    def _parse_pasal_structure(self, text):
        """Parse text into Pasal-Ayat-Huruf structure with Penjelasan support"""
        lines = text.splitlines()
        pasal_blocks = []
        toc = []
        
        current_pasal = None
        current_pasal_data = None
        current_ayat = None
        in_penjelasan_section = False
        penjelasan_ayat_map = {}
        
        # Regex patterns
        pasal_pat = re.compile(r'^Pasal\s+(\d+)')
        ayat_pat = re.compile(r'^\((\d+)\)\s*(.*)')
        huruf_pat = re.compile(r'^\s{4}([a-z])\.\s*(.*)')
        penjelasan_section_pat = re.compile(r'^Penjelasan\s*$')
        penjelasan_pasal_pat = re.compile(r'^Pasal\s+(\d+)\s*$')
        penjelasan_ayat_pat = re.compile(r'^Penjelasan\s+\((\d+)\):\s*(.*)')
        penjelasan_huruf_pat = re.compile(r'^Penjelasan\s+\((\d+)\)([a-z]):\s*(.*)')
        bab_pat = re.compile(r'^BAB\s+[IVX]+')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            if bab_pat.match(line):
                i += 1
                continue
            
            if penjelasan_section_pat.match(line):
                in_penjelasan_section = True
                i += 1
                continue
            
            if in_penjelasan_section:
                pasal_penjelasan_match = penjelasan_pasal_pat.match(line)
                if pasal_penjelasan_match:
                    current_pasal_penjelasan = pasal_penjelasan_match.group(1)
                    i += 1
                    temp_penjelasan = []
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if not next_line:
                            i += 1
                            continue
                        if penjelasan_ayat_pat.match(next_line) or penjelasan_pasal_pat.match(next_line) or pasal_pat.match(next_line):
                            break
                        temp_penjelasan.append(next_line)
                        i += 1
                    
                    if temp_penjelasan:
                        penjelasan_text = ' '.join(temp_penjelasan)
                        if current_pasal_penjelasan not in penjelasan_ayat_map:
                            penjelasan_ayat_map[current_pasal_penjelasan] = {'umum': penjelasan_text, 'ayat': {}}
                        else:
                            penjelasan_ayat_map[current_pasal_penjelasan]['umum'] = penjelasan_text
                    continue
                
                penjelasan_ayat_match = penjelasan_ayat_pat.match(line)
                if penjelasan_ayat_match and current_pasal:
                    ayat_num = penjelasan_ayat_match.group(1)
                    penjelasan_text = penjelasan_ayat_match.group(2).strip()
                    
                    if current_pasal not in penjelasan_ayat_map:
                        penjelasan_ayat_map[current_pasal] = {'umum': '', 'ayat': {}}
                    penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = penjelasan_text
                    i += 1
                    continue
                
                penjelasan_huruf_match = penjelasan_huruf_pat.match(line)
                if penjelasan_huruf_match and current_pasal:
                    ayat_num = penjelasan_huruf_match.group(1)
                    huruf = penjelasan_huruf_match.group(2)
                    penjelasan_text = penjelasan_huruf_match.group(3).strip()
                    
                    if current_pasal not in penjelasan_ayat_map:
                        penjelasan_ayat_map[current_pasal] = {'umum': '', 'ayat': {}}
                    if ayat_num not in penjelasan_ayat_map[current_pasal]['ayat']:
                        penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = {}
                    if not isinstance(penjelasan_ayat_map[current_pasal]['ayat'][ayat_num], dict):
                        penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = {'_text': penjelasan_ayat_map[current_pasal]['ayat'][ayat_num]}
                    
                    penjelasan_ayat_map[current_pasal]['ayat'][ayat_num][huruf] = penjelasan_text
                    i += 1
                    continue
            
            pasal_match = pasal_pat.match(line)
            if pasal_match and not in_penjelasan_section:
                if current_pasal_data:
                    pasal_blocks.append(current_pasal_data)
                    toc.append({'pasal': current_pasal, 'ayats': [a['nomor'] for a in current_pasal_data['ayats']]})
                
                current_pasal = pasal_match.group(1)
                current_pasal_data = {'nomor': current_pasal, 'ayats': [], 'penjelasan_umum': ''}
                current_ayat = None
                i += 1
                continue
            
            ayat_match = ayat_pat.match(line)
            if ayat_match and current_pasal_data and not in_penjelasan_section:
                ayat_num = ayat_match.group(1)
                ayat_text = ayat_match.group(2).strip()
                
                current_ayat = {'nomor': ayat_num, 'isi': [ayat_text] if ayat_text else [], 'hurufs': [], 'penjelasan': ''}
                current_pasal_data['ayats'].append(current_ayat)
                i += 1
                continue
            
            huruf_match = huruf_pat.match(line)
            if huruf_match and current_ayat and not in_penjelasan_section:
                huruf = huruf_match.group(1)
                huruf_text = huruf_match.group(2).strip()
                current_ayat['hurufs'].append({'huruf': huruf, 'isi': huruf_text, 'penjelasan': ''})
                i += 1
                continue
            
            if current_ayat and not in_penjelasan_section and line:
                current_ayat['isi'].append(line)
            
            i += 1
        
        if current_pasal_data:
            pasal_blocks.append(current_pasal_data)
            toc.append({'pasal': current_pasal, 'ayats': [a['nomor'] for a in current_pasal_data['ayats']]})
        
        # Merge penjelasan
        for pasal in pasal_blocks:
            pasal_num = pasal['nomor']
            if pasal_num in penjelasan_ayat_map:
                pasal['penjelasan_umum'] = penjelasan_ayat_map[pasal_num].get('umum', '')
                
                for ayat in pasal['ayats']:
                    ayat_num = ayat['nomor']
                    if ayat_num in penjelasan_ayat_map[pasal_num].get('ayat', {}):
                        penjelasan_data = penjelasan_ayat_map[pasal_num]['ayat'][ayat_num]
                        
                        if isinstance(penjelasan_data, str):
                            ayat['penjelasan'] = penjelasan_data
                        elif isinstance(penjelasan_data, dict):
                            ayat['penjelasan'] = penjelasan_data.get('_text', '')
                            
                            for huruf_item in ayat['hurufs']:
                                huruf = huruf_item['huruf']
                                if huruf in penjelasan_data:
                                    huruf_item['penjelasan'] = penjelasan_data[huruf]
        
        return pasal_blocks, toc

    def _generate_structured_html(self, pasal_blocks, toc):
        """Generate HTML with accordion + TOC + Penjelasan"""
        html_parts = []
        
        # CSS
        html_parts.append('''
<style>
.regulation-container { display: flex; gap: 20px; }
.toc-sidebar { flex: 0 0 250px; position: sticky; top: 20px; max-height: 80vh; overflow-y: auto; 
               padding: 15px; background: #f8f9fa; border-radius: 8px; border: 1px solid #dee2e6; }
.toc-sidebar h5 { margin-bottom: 15px; color: #495057; font-weight: 600; font-size: 1em; }
.toc-sidebar ul { list-style: none; padding-left: 0; margin: 0; }
.toc-sidebar ul ul { padding-left: 15px; margin-top: 5px; }
.toc-sidebar a { color: #007bff; text-decoration: none; display: block; padding: 3px 0; font-size: 0.9em; }
.toc-sidebar a:hover { color: #0056b3; text-decoration: underline; }
.regulation-content { flex: 1; }
.pasal-card { margin-bottom: 20px; border: 1px solid #dee2e6; border-radius: 8px; overflow: hidden; }
.pasal-header { background: #007bff; color: white; padding: 12px 15px; cursor: pointer; 
                font-weight: 600; font-size: 1.05em; }
.pasal-header:hover { background: #0056b3; }
.pasal-content { padding: 20px; background: white; }
.ayat { margin-bottom: 15px; padding: 12px; background: #f8f9fa; border-left: 4px solid #007bff; border-radius: 4px; }
.ayat strong { color: #007bff; font-size: 1.05em; }
.ayat ul { margin-top: 8px; padding-left: 20px; list-style: none; }
.ayat li { margin-bottom: 8px; }
.penjelasan { margin-top: 10px; padding: 10px; background: #fff3cd; border-left: 4px solid #ffc107; 
              border-radius: 4px; font-style: italic; font-size: 0.95em; }
.penjelasan strong { color: #856404; font-style: normal; }
.penjelasan-umum { margin-bottom: 15px; padding: 12px; background: #d1ecf1; border-left: 4px solid #17a2b8; 
                   border-radius: 4px; font-size: 0.95em; }
.penjelasan-umum strong { color: #0c5460; }
</style>
''')
        
        html_parts.append('<div class="regulation-container">')
        
        # TOC Sidebar
        html_parts.append('<div class="toc-sidebar"><h5>📋 Daftar Isi</h5><ul>')
        for entry in toc:
            html_parts.append(f'<li><a href="#pasal{entry["pasal"]}">Pasal {entry["pasal"]}</a>')
            if entry['ayats']:
                html_parts.append('<ul>')
                for ayat in entry['ayats']:
                    html_parts.append(f'<li><a href="#pasal{entry["pasal"]}-ayat{ayat}">Ayat ({ayat})</a></li>')
                html_parts.append('</ul>')
            html_parts.append('</li>')
        html_parts.append('</ul></div>')
        
        # Content
        html_parts.append('<div class="regulation-content">')
        
        for pasal in pasal_blocks:
            pasal_id = f'pasal{pasal["nomor"]}'
            html_parts.append(f'<div class="pasal-card" id="{pasal_id}">')
            html_parts.append(f'<div class="pasal-header" onclick="this.nextElementSibling.style.display=this.nextElementSibling.style.display===\'none\'?\'block\':\'none\'">📜 Pasal {pasal["nomor"]}</div>')
            html_parts.append(f'<div class="pasal-content" style="display:block;">')
            
            if pasal.get('penjelasan_umum'):
                html_parts.append(f'<div class="penjelasan-umum"><strong>💡 Penjelasan Pasal:</strong><br>{pasal["penjelasan_umum"]}</div>')
            
            for ayat in pasal['ayats']:
                ayat_id = f'{pasal_id}-ayat{ayat["nomor"]}'
                ayat_text = ' '.join(ayat['isi']).strip()
                html_parts.append(f'<div class="ayat" id="{ayat_id}"><strong>({ayat["nomor"]})</strong> {ayat_text}')
                
                if ayat['hurufs']:
                    html_parts.append('<ul>')
                    for huruf in ayat['hurufs']:
                        html_parts.append(f'<li><strong>{huruf["huruf"]}.</strong> {huruf["isi"]}')
                        if huruf.get('penjelasan'):
                            html_parts.append(f'<div class="penjelasan"><strong>💡 Penjelasan huruf {huruf["huruf"]}:</strong> {huruf["penjelasan"]}</div>')
                        html_parts.append('</li>')
                    html_parts.append('</ul>')
                
                if ayat.get('penjelasan'):
                    html_parts.append(f'<div class="penjelasan"><strong>💡 Penjelasan ayat ({ayat["nomor"]}):</strong> {ayat["penjelasan"]}</div>')
                
                html_parts.append('</div>')
            
            html_parts.append('</div></div>')
        
        html_parts.append('</div></div>')
        
        return '\n'.join(html_parts)

    def _insert_penjelasan_into_text(self, text):
        """
        Insert Penjelasan sections into their respective Pasal/Ayat/Huruf positions.
        
        Handles complex hierarchical structure:
        - KUBU 1 (Main Content): Regular regulation text with Pasal, Ayat, Huruf
        - KUBU 2 (Penjelasan): Starts with "PENJELASAN ATAS" or "Penjelasan"
        
        Format in Penjelasan section:
        - Pasal X (standalone) = explanation for whole Pasal
        - Ayat (X) (under context of Pasal) = explanation for specific Ayat
        - Huruf X (under context of Pasal and Ayat) = explanation for specific Huruf
        """
        try:
            lines = text.split('\n')
            result_lines = []
            penjelasan_map = {}  # key: "pasal_X", "pasal_X_ayat_Y", "pasal_X_ayat_Y_huruf_Z", "umum"
            in_penjelasan_section = False
            current_penjelasan_context = {'pasal': None, 'ayat': None}
            current_penjelasan_key = None
            current_penjelasan_buffer = []
            
            # Regex patterns for detecting Penjelasan section
            penjelasan_section_pat = re.compile(r'^(PENJELASAN\s+ATAS|PENJELASAN|Penjelasan)\s*$', re.IGNORECASE)
            
            # Pattern for "I. UMUM" or variations
            umum_pattern = re.compile(r'^\s*I\.?\s*UMUM\s*$', re.IGNORECASE)
            pasal_demi_pasal_pattern = re.compile(r'^\s*II\.?\s*PASAL\s+DEMI\s+PASAL\s*$', re.IGNORECASE)
            
            # Patterns in Penjelasan section
            pasal_in_penjelasan_pat = re.compile(r'^Pasal\s+(\d+[A-Z]?)\s*$', re.IGNORECASE)  # Match "Pasal 1" or "Pasal 45A"
            # SPECIAL: Also match Pasal with 3 digits (for OCR errors like "Pasal 458")
            pasal_3digit_pat = re.compile(r'^Pasal\s+(\d{3})\s*$', re.IGNORECASE)  # Match "Pasal 458", "Pasal 457"
            # ALTERNATIVE: Match "Pasal 1." (with period) - old format - can have content after period on same line
            pasal_with_period_pat = re.compile(r'^Pasal\s+(\d+[A-Z]?)\.\s*(.*)$', re.IGNORECASE)  # Match "Pasal 1." or "Pasal 1. content"
            
            # Pattern to detect multi-pasal format (e.g., "Pasal 5, 6 dan 7.") - NOT a header, should be skipped
            multi_pasal_pat = re.compile(r'^Pasal\s+[\d,\s]+dan\s+\d+\s*\.?\s*$', re.IGNORECASE)
            
            # FIXED: Support "Ayat (1)", "Ayat 1", "Ayat (2a)", "Ayat 2b" - with optional letters
            ayat_header_pat = re.compile(r'^Ayat\s+\(?(\d+[a-z]?)\)?\s*$', re.IGNORECASE)
            # ALTERNATIVE: Match "Ayat 1." (with period) - old format - can have content after period on same line
            ayat_with_period_pat = re.compile(r'^Ayat\s+(\d+[a-z]?)\.\s*(.*)$', re.IGNORECASE)  # Match "Ayat 1." or "Ayat 1. content"
            
            # FIXED: Match "Huruf x" exactly (standalone, not inline with text)
            # Also support variations like "Huruf j dan Huruf k" or "Huruf j dan k"
            huruf_header_pat = re.compile(r'^Huruf\s+([a-z])\s*$', re.IGNORECASE)
            # ALTERNATIVE: Match "Huruf x." (with period) - old format - can have content after period on same line
            huruf_with_period_pat = re.compile(r'^Huruf\s+([a-z])\.\s*(.*)$', re.IGNORECASE)  # Match "Huruf a." or "Huruf a. content"
            huruf_combined_pat = re.compile(r'^Huruf\s+([a-z])\s+dan\s+(?:Huruf\s+)?([a-z])\s*$', re.IGNORECASE)  # "Huruf j dan k"
            
            # Pattern to detect "Cukup jelas" (means no explanation needed)
            cukup_jelas_pat = re.compile(r'^Cukup\s+jelas\.?\s*$', re.IGNORECASE)
            # Pattern to detect "Cukup jelas" anywhere in text (including inline)
            cukup_jelas_inline_pat = re.compile(r'Cukup\s+jelas\.?', re.IGNORECASE)
            
            # Helper function to clean content - remove everything after "Cukup jelas"
            def clean_cukup_jelas(content):
                """Remove all text after 'Cukup jelas' if found, and truncate at last period before 'TAMBAHAN LEMBARAN NEGARA'"""
                if not content:
                    return None
                
                # Log untuk debugging
                original_content = content
                content_lower = content.lower()
                
                # Check if content is ONLY "Cukup jelas" (with optional trailing period/whitespace)
                if re.match(r'^\s*cukup\s+jelas\.?\s*$', content, re.IGNORECASE):
                    _logger.info(f"  [NO] Content is only 'Cukup jelas' - discarding")
                    return None
                
                # First, check for "TAMBAHAN LEMBARAN NEGARA" and truncate at last period before it
                tambahan_match = re.search(r'TAMBAHAN\s+LEMBARAN\s+NEGARA', content, re.IGNORECASE)
                if tambahan_match:
                    # Find the last period before "TAMBAHAN LEMBARAN NEGARA"
                    content_before_tambahan = content[:tambahan_match.start()]
                    last_period_idx = content_before_tambahan.rfind('.')
                    if last_period_idx > 0:
                        content = content[:last_period_idx + 1].strip()
                        _logger.info(f"  [CUT] Truncated at last period before 'TAMBAHAN LEMBARAN NEGARA'")
                
                # Check if "Cukup jelas" appears in content - cut everything after it
                # Use case-insensitive search
                match = re.search(r'cukup\s+jelas\.?', content, re.IGNORECASE)
                if match:
                    # Found "Cukup jelas" - truncate content before it
                    before_cukup = content[:match.start()].strip()
                    
                    # Remove trailing punctuation/whitespace
                    before_cukup = before_cukup.rstrip('.,;: \t\n')
                    
                    # Clean up any lines that only have period (.) left
                    lines_cleaned = [l for l in before_cukup.split('\n') if l.strip() and l.strip() != '.']
                    before_cukup = '\n'.join(lines_cleaned).strip()
                    
                    _logger.info(f"  🔪 Found 'Cukup jelas' at position {match.start()} - truncating")
                    _logger.info(f"     Original: '{content[:100]}...'")
                    _logger.info(f"     Truncated to: '{before_cukup[:100]}...'")
                    
                    # If nothing meaningful before "Cukup jelas", return None
                    if len(before_cukup) < 10:
                        _logger.info(f"  [NO] Content before 'Cukup jelas' too short ({len(before_cukup)} chars) - discarding")
                        return None
                    
                    return before_cukup
                
                # Check for "PRESIDEN REPUBLIK INDONESIA-X" pattern (footer/header leak)
                if re.match(r'^PRESIDEN\s+REPUBLIK\s+INDONESIA-\d+', content.strip(), re.IGNORECASE):
                    _logger.info(f"  [NO] Found 'PRESIDEN REPUBLIK INDONESIA-X' footer - discarding: '{content[:80]}...'")
                    return None
                
                return content
            
            # Pattern to detect continuation markers (e.g., "Huruf d ...", "Pasal 2 ...")
            continuation_marker_pat = re.compile(r'^(Pasal|Ayat|Huruf)\s+.*\.\.\.\s*$', re.IGNORECASE)
            continuation_marker_pat = re.compile(r'^(Pasal|Ayat|Huruf)\s+.*\.\.\.\s*$', re.IGNORECASE)
            
            _logger.info("=" * 80)
            _logger.info("PARSING PENJELASAN SECTION")
            _logger.info("=" * 80)
            
            # FIRST PASS: Collect all penjelasan entries
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Detect start of Penjelasan section
                if penjelasan_section_pat.match(line):
                    in_penjelasan_section = True
                    _logger.info(f"[OK] Found Penjelasan section at line {i}: {line}")
                    i += 1
                    continue
                
                if in_penjelasan_section:
                    # Check for "I. UMUM" section
                    if umum_pattern.match(line):
                        # Save previous penjelasan if any
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            content = clean_cukup_jelas(content)
                            if content:
                                penjelasan_map[current_penjelasan_key] = content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {content[:50]}...")
                        
                        # Set context to UMUM
                        current_penjelasan_key = 'umum'
                        current_penjelasan_context = {'pasal': None, 'ayat': None}
                        current_penjelasan_buffer = []
                        _logger.info(f"\n-> Found 'I. UMUM' - collecting Penjelasan Umum content")
                        i += 1
                        continue
                    
                    # Check for "II. PASAL DEMI PASAL" - end of UMUM section
                    if pasal_demi_pasal_pattern.match(line):
                        # Save UMUM content if we were in it
                        if current_penjelasan_key == 'umum' and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            content = clean_cukup_jelas(content)
                            if content and len(content) > 10:
                                penjelasan_map['umum'] = content
                                _logger.info(f"  Saved [umum]: {content[:80]}...")
                        
                        # Reset context for Pasal Demi Pasal
                        current_penjelasan_key = None
                        current_penjelasan_buffer = []
                        _logger.info(f"\n-> Found 'II. PASAL DEMI PASAL' - ending Penjelasan Umum")
                        i += 1
                        continue
                    
                    # IMPORTANT: Skip multi-pasal format (e.g., "Pasal 5, 6 dan 7.") - this is NOT a header
                    # Also skip its content (usually "Sudah cukup terang, tidak memperlukan penjelasan")
                    if multi_pasal_pat.match(line):
                        _logger.info(f"  [SKIP] Skipping multi-pasal format: {line.strip()}")
                        
                        # First, save current buffer (for previous Pasal) before skipping
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            content = clean_cukup_jelas(content)
                            if content:
                                penjelasan_map[current_penjelasan_key] = content
                                _logger.info(f"  Saved [{current_penjelasan_key}] before skip: {content[:50]}...")
                        
                        # Reset context and buffer
                        current_penjelasan_context = {'pasal': None, 'ayat': None}
                        current_penjelasan_key = None
                        current_penjelasan_buffer = []
                        
                        # Skip content lines until we find next empty line or next Pasal/section
                        i += 1
                        while i < len(lines):
                            skip_line = lines[i].strip()
                            if not skip_line:
                                # Empty line - stop skipping
                                break
                            # Check if next line is a valid Pasal header (not multi-pasal)
                            if (pasal_in_penjelasan_pat.match(skip_line) or 
                                pasal_3digit_pat.match(skip_line) or 
                                (pasal_with_period_pat.match(skip_line) and not multi_pasal_pat.match(skip_line))):
                                # Found next Pasal - don't consume this line, let normal flow handle it
                                i -= 1  # Step back so next iteration processes this line
                                break
                            _logger.info(f"    ↳ Skipping content: {skip_line[:60]}...")
                            i += 1
                        
                        i += 1
                        continue
                    
                    # Check for Pasal declaration (standalone = explanation for whole Pasal)
                    pasal_match = pasal_in_penjelasan_pat.match(line)
                    pasal_3digit_match = pasal_3digit_pat.match(line) if not pasal_match else None
                    pasal_period_match = pasal_with_period_pat.match(line) if not pasal_match and not pasal_3digit_match else None
                    
                    if pasal_match or pasal_3digit_match or pasal_period_match:
                        # Save previous penjelasan if any
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            content = clean_cukup_jelas(content)
                            if content:
                                penjelasan_map[current_penjelasan_key] = content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {content[:50]}...")
                        
                        # Set new context
                        if pasal_match:
                            pasal_num = pasal_match.group(1)
                        elif pasal_3digit_match:
                            pasal_num = pasal_3digit_match.group(1)
                        else:
                            pasal_num = pasal_period_match.group(1)
                        
                        # FIX OCR errors in Pasal number DURING PENJELASAN PARSING
                        # This ensures "Pasal 458" in PENJELASAN section is recognized as "45B"
                        if pasal_num == '458':
                            pasal_num = '45B'
                            _logger.info(f"  [OK] OCR fix: Pasal 458 -> Pasal 45B in PENJELASAN section")
                        elif pasal_num == '457':
                            pasal_num = '45A'
                            _logger.info(f"  [OK] OCR fix: Pasal 457 -> Pasal 45A in PENJELASAN section")
                        
                        current_penjelasan_context = {'pasal': pasal_num, 'ayat': None}
                        
                        # Peek ahead: check if next non-empty line is "Ayat" or "Huruf"
                        # If yes, don't set penjelasan_key to pasal (no pasal-level explanation)
                        peek_ahead_idx = i + 1
                        has_immediate_ayat_or_huruf = False
                        while peek_ahead_idx < len(lines):
                            peek_line = lines[peek_ahead_idx].strip()
                            if peek_line:
                                # Check if it's an Ayat or Huruf header (both old and new format)
                                if (ayat_header_pat.match(peek_line) or 
                                    ayat_with_period_pat.match(peek_line) or
                                    huruf_header_pat.match(peek_line)):
                                    has_immediate_ayat_or_huruf = True
                                break
                            peek_ahead_idx += 1
                        
                        if has_immediate_ayat_or_huruf:
                            # Don't set pasal key - pasal has no direct explanation
                            current_penjelasan_key = None
                            _logger.info(f"\n-> Pasal {pasal_num} context started (no pasal-level explanation, has Ayat/Huruf)")
                        else:
                            # Set pasal key - pasal has explanation
                            current_penjelasan_key = f'pasal_{pasal_num}'
                            _logger.info(f"\n-> Pasal {pasal_num} context started (has pasal-level explanation)")
                        
                        current_penjelasan_buffer = []
                        
                        # If Pasal has content on same line (old format "Pasal 1. content"), add it to buffer
                        if pasal_period_match and pasal_period_match.group(2).strip():
                            current_penjelasan_buffer.append(pasal_period_match.group(2).strip())
                        
                        i += 1
                        continue
                    
                    # Check for Ayat header (e.g., "Ayat (1)" or "Ayat 1.")
                    ayat_match = ayat_header_pat.match(line)
                    ayat_period_match = ayat_with_period_pat.match(line) if not ayat_match else None
                    
                    if (ayat_match or ayat_period_match) and current_penjelasan_context['pasal']:
                        ayat_num = ayat_match.group(1) if ayat_match else ayat_period_match.group(1)
                        
                        # CRITICAL FIX: Detect implicit Pasal boundary
                        # If we see "Ayat (1)" again, it means we've moved to a new Pasal
                        # (even if there's no "Pasal X" header)
                        current_ayat = current_penjelasan_context.get('ayat')
                        if ayat_num == '1' and current_ayat and current_ayat != '1':
                            # We found "Ayat (1)" but we already processed other ayats
                            # This means implicit new Pasal - STOP collecting for previous Pasal
                            _logger.warning(f"  [WARN] IMPLICIT PASAL BOUNDARY DETECTED!")
                            _logger.warning(f"     Found 'Ayat (1)' but current context is Pasal {current_penjelasan_context['pasal']} Ayat ({current_ayat})")
                            _logger.warning(f"     This indicates start of NEW Pasal without header")
                            _logger.warning(f"     DISCARDING current buffer to prevent cross-contamination")
                            
                            # Clear context and buffer - we don't know which Pasal this belongs to
                            current_penjelasan_context = {'pasal': None, 'ayat': None}
                            current_penjelasan_key = None
                            current_penjelasan_buffer = []
                            i += 1
                            continue
                        
                        # Save previous penjelasan (could be for Pasal or previous Ayat)
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            # Check if content is meaningful (not just "Ayat X Cukup jelas" repeated)
                            # Filter out buffer that only contains headers like "Ayat 1", "Ayat 1.", "Ayat 2", "Ayat 2b", etc
                            # Also filter out lines with only period (.) left from header removal
                            filtered_lines = [
                                line for line in current_penjelasan_buffer 
                                if line.strip() and 
                                line.strip() != '.' and  # Remove lonely period
                                not re.match(r'^Ayat\s+\(?\d+[a-z]?\)?\s*\.?\s*$', line.strip()) and  # Support both "Ayat 1" and "Ayat 1."
                                not re.match(r'^Huruf\s+[a-z]\s*\.?\s*$', line.strip()) and  # Support both "Huruf a" and "Huruf a."
                                not re.match(r'^Pasal\s+\d+[A-Z]?\s*\.?\s*$', line.strip()) and  # Filter out "Pasal X" or "Pasal X."
                                not re.match(r'^Pasal\s+[\d,\s]+dan\s+\d+\s*\.?\s*$', line.strip(), re.IGNORECASE)  # Filter "Pasal 5, 6 dan 7."
                            ]
                            # Preserve newlines for list items (a. b. c.) but join paragraphs
                            filtered_content = '\n'.join(filtered_lines).strip()
                            
                            filtered_content = clean_cukup_jelas(filtered_content)
                            if filtered_content and len(filtered_content) > 10:
                                # Check for duplicate key (Ayat section)
                                if current_penjelasan_key in penjelasan_map:
                                    existing_preview = penjelasan_map[current_penjelasan_key][:80]
                                    new_preview = filtered_content[:80]
                                    _logger.warning(f"  [WARN] DUPLICATE KEY DETECTED: [{current_penjelasan_key}]")
                                    _logger.warning(f"     FIRST (kept): {existing_preview}...")
                                    _logger.warning(f"     DUPLICATE (ignored): {new_preview}...")
                                    _logger.warning(f"     Keeping FIRST occurrence, ignoring duplicate Ayat header")
                                else:
                                    penjelasan_map[current_penjelasan_key] = filtered_content
                                    _logger.info(f"  Saved [{current_penjelasan_key}]: {filtered_content[:50]}...")
                        
                        # Set new Ayat context
                        ayat_num = ayat_match.group(1) if ayat_match else ayat_period_match.group(1)
                        current_penjelasan_context['ayat'] = ayat_num
                        pasal_num = current_penjelasan_context['pasal']
                        
                        # Peek ahead: check if next non-empty line is "Huruf"
                        peek_ahead_idx = i + 1
                        has_immediate_huruf = False
                        while peek_ahead_idx < len(lines):
                            peek_line = lines[peek_ahead_idx].strip()
                            if peek_line:
                                # Check both old and new Huruf format
                                if huruf_header_pat.match(peek_line) or huruf_with_period_pat.match(peek_line):
                                    has_immediate_huruf = True
                                break
                            peek_ahead_idx += 1
                        
                        if has_immediate_huruf:
                            # Don't set ayat key - ayat has no direct explanation, has Huruf
                            current_penjelasan_key = None
                            _logger.info(f"  -> Ayat ({ayat_num}) context started (no ayat-level explanation, has Huruf)")
                        else:
                            # Set ayat key - ayat has explanation
                            current_penjelasan_key = f"pasal_{pasal_num}_ayat_{ayat_num}"
                            _logger.info(f"  -> Ayat ({ayat_num}) context started (has ayat-level explanation)")
                        
                        current_penjelasan_buffer = []
                        
                        # If Ayat has content on same line (old format "Ayat 1. content"), add it to buffer
                        if ayat_period_match and ayat_period_match.group(2).strip():
                            current_penjelasan_buffer.append(ayat_period_match.group(2).strip())
                        
                        i += 1
                        continue
                    
                    # Check for Huruf header (e.g., "Huruf a" or "Huruf a.")
                    huruf_match = huruf_header_pat.match(line)
                    huruf_period_match = huruf_with_period_pat.match(line) if not huruf_match else None
                    huruf_combined_match = huruf_combined_pat.match(line) if not huruf_match and not huruf_period_match else None
                    
                    if (huruf_match or huruf_period_match or huruf_combined_match) and current_penjelasan_context['pasal'] and current_penjelasan_context['ayat']:
                        # Save previous penjelasan
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = '\n'.join(current_penjelasan_buffer).strip()
                            # Filter out headers and preserve newlines for list items
                            # Also filter out lines with only period (.) left from header removal
                            filtered_lines = [
                                line for line in current_penjelasan_buffer 
                                if line.strip() and 
                                line.strip() != '.' and  # Remove lonely period
                                not re.match(r'^Ayat\s+\(?\d+[a-z]?\)?\s*\.?\s*$', line.strip()) and  # Support both "Ayat 1" and "Ayat 1."
                                not re.match(r'^Huruf\s+[a-z]\s*\.?\s*$', line.strip()) and  # Support both "Huruf a" and "Huruf a."
                                not re.match(r'^Pasal\s+\d+[A-Z]?\s*\.?\s*$', line.strip()) and  # Filter out "Pasal X" or "Pasal X."
                                not re.match(r'^Pasal\s+[\d,\s]+dan\s+\d+\s*\.?\s*$', line.strip(), re.IGNORECASE)  # Filter "Pasal 5, 6 dan 7."
                            ]
                            filtered_content = '\n'.join(filtered_lines).strip()
                            
                            filtered_content = clean_cukup_jelas(filtered_content)
                            if filtered_content and len(filtered_content) > 10:
                                penjelasan_map[current_penjelasan_key] = filtered_content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {filtered_content[:50]}...")
                                
                                # If previous was a combined huruf, save to BOTH
                                if 'huruf_combined' in current_penjelasan_context:
                                    huruf1, huruf2 = current_penjelasan_context['huruf_combined']
                                    pasal_num = current_penjelasan_context['pasal']
                                    ayat_num = current_penjelasan_context['ayat']
                                    key2 = f"pasal_{pasal_num}_ayat_{ayat_num}_huruf_{huruf2}"
                                    penjelasan_map[key2] = filtered_content
                                    _logger.info(f"  Also saved to [{key2}]: {filtered_content[:50]}... (combined huruf)")
                            else:
                                _logger.info(f"  [WARN] Skipped [{current_penjelasan_key}]: content too short ({len(filtered_content) if filtered_content else 0} chars) after cleaning")
                        
                        pasal_num = current_penjelasan_context['pasal']
                        ayat_num = current_penjelasan_context['ayat']
                        
                        if huruf_combined_match:
                            # "Huruf j dan k" - save for both j and k
                            huruf1 = huruf_combined_match.group(1)
                            huruf2 = huruf_combined_match.group(2)
                            _logger.info(f"    -> Combined Huruf {huruf1} dan {huruf2} context started (Pasal {pasal_num} Ayat {ayat_num})")
                            # We'll collect the content and save it for BOTH huruf
                            current_penjelasan_key = f"pasal_{pasal_num}_ayat_{ayat_num}_huruf_{huruf1}"  # Start with first
                            current_penjelasan_context['huruf_combined'] = (huruf1, huruf2)  # Track both
                        else:
                            # Single huruf (either "Huruf a" or "Huruf a.")
                            huruf = huruf_match.group(1) if huruf_match else huruf_period_match.group(1)
                            current_penjelasan_key = f"pasal_{pasal_num}_ayat_{ayat_num}_huruf_{huruf}"
                            current_penjelasan_context.pop('huruf_combined', None)  # Clear combined flag
                            _logger.info(f"    -> Huruf {huruf} context started (Pasal {pasal_num} Ayat {ayat_num})")
                        
                        current_penjelasan_buffer = []
                        
                        # If Huruf has content on same line (old format "Huruf a. content"), add it to buffer
                        if huruf_period_match and huruf_period_match.group(2).strip():
                            current_penjelasan_buffer.append(huruf_period_match.group(2).strip())
                        
                        i += 1
                        continue
                    
                    # Accumulate text for current penjelasan (skip empty lines at start)
                    if line:
                        # Check if this is "Cukup jelas" - if so, skip adding to buffer and mark it
                        if cukup_jelas_pat.match(line):
                            # Clear buffer - we don't want to save "Cukup jelas" entries
                            current_penjelasan_buffer = []
                            # Set key to None to prevent saving
                            current_penjelasan_key = None
                            _logger.info(f"  [WARN] Found 'Cukup jelas' - clearing buffer and key")
                            i += 1
                            continue
                        
                        # If key is None (we're after "Cukup jelas"), don't accumulate anything
                        # until we hit a new Pasal/Ayat/Huruf header
                        if current_penjelasan_key is None:
                            i += 1
                            continue
                        
                        # Skip continuation markers, page markers, and headers
                        if (not continuation_marker_pat.match(line) and
                            not re.match(r'^\s*-\s*\d+\s*-\s*$', line) and  # Page markers
                            not re.match(r'^Ayat\s+\(?\d+[a-z]?\)?\s*$', line.strip()) and  # Ayat headers: "Ayat (1)", "Ayat 2b", etc
                            not re.match(r'^Huruf\s+[a-z]\s*$', line.strip())):  # Huruf headers
                            if current_penjelasan_buffer or line.strip():  # Skip leading empty lines
                                current_penjelasan_buffer.append(line)
                
                i += 1
            
            # Save last penjelasan (only if content is not empty and not just "Cukup jelas")
            if current_penjelasan_key and current_penjelasan_buffer:
                content = '\n'.join(current_penjelasan_buffer).strip()
                content = clean_cukup_jelas(content)
                # Double check: content should have meaningful text
                if content and len(content) > 10:  # More than just whitespace or short placeholder
                    penjelasan_map[current_penjelasan_key] = content
                    _logger.info(f"  Saved [{current_penjelasan_key}]: {content[:50]}...")
                    
                    # If this was a combined huruf (j dan k), save to BOTH huruf
                    if 'huruf_combined' in current_penjelasan_context:
                        huruf1, huruf2 = current_penjelasan_context['huruf_combined']
                        pasal_num = current_penjelasan_context['pasal']
                        ayat_num = current_penjelasan_context['ayat']
                        key2 = f"pasal_{pasal_num}_ayat_{ayat_num}_huruf_{huruf2}"
                        penjelasan_map[key2] = content
                        _logger.info(f"  Also saved to [{key2}]: {content[:50]}... (combined huruf)")

            
            _logger.info(f"\n[OK] Total penjelasan collected: {len(penjelasan_map)}")
            for key in sorted(penjelasan_map.keys()):
                _logger.info(f"  - {key}: {penjelasan_map[key][:60]}...")
            
            # SECOND PASS: Insert penjelasan into main content
            in_penjelasan_section = False
            current_pasal = None
            current_ayat = None
            umum_inserted = False  # Track if we've inserted Penjelasan Umum
            
            # pending penjelasan markers to flush at boundaries
            pending_pasal_lines = None
            pending_ayat_lines = None
            pending_huruf_lines = None
            
            # Patterns for main content
            pasal_main_pat = re.compile(r'^Pasal\s+((?:\d+[A-Z]?)|(?:[IVX]+))', re.IGNORECASE)  # Match "Pasal 1", "Pasal 45A", or "Pasal II"
            ayat_main_pat = re.compile(r'^\((\d+)\)')
            huruf_main_pat = re.compile(r'^([a-z])\.')
            bab_header_pat = re.compile(r'^(BAB|BAGIAN)\s+', re.IGNORECASE)
            # Pattern to detect "Menimbang" (where we insert Penjelasan Umum before it)
            # Support both with and without colon
            menimbang_pat = re.compile(r'^\s*(Menimbang)\s*:?\s*(.*)$', re.IGNORECASE)
            # Pattern to detect "DENGAN RAHMAT ... PRESIDEN REPUBLIK INDONESIA" - insert Penjelasan Umum AFTER it
            dengan_rahmat_presiden_pat = re.compile(r'^\s*DENGAN\s+RAHMAT.*PRESIDEN\s+REPUBLIK\s+INDONESIA', re.IGNORECASE)
            
            _logger.info("\n" + "=" * 80)
            _logger.info("INSERTING PENJELASAN INTO MAIN CONTENT (AFTER CONTENT)")
            _logger.info("=" * 80)
            
            def render_penjelasan_text(txt: str) -> str:
                # Render enumerations ((1), 1., a.) into ordered lists
                # Support multiple lists separated by intro text like "seperti berikut:"
                if not txt:
                    return ''
                # Normalize newlines but PRESERVE empty lines to detect list boundaries
                txt_n = txt.replace('\r\n', '\n').replace('\r', '\n')
                lines = txt_n.split('\n')  # Don't strip empty lines yet
                
                # Pattern to detect list introduction (e.g., "seperti berikut:", "yaitu:", "sebagai berikut:")
                list_intro_pat = re.compile(r'.*(seperti berikut|sebagai berikut|yaitu|antara lain)\s*:\s*$', re.IGNORECASE)
                
                html_parts = []
                current_list_items = []
                current_list_type = None
                current_paragraph_lines = []
                prev_was_empty = False
                
                def flush_paragraph():
                    """Flush accumulated paragraph lines"""
                    nonlocal current_paragraph_lines
                    if current_paragraph_lines:
                        para_text = ' '.join(current_paragraph_lines)
                        html_parts.append(f'<p style="margin-bottom: 10px;">{para_text}</p>')
                        current_paragraph_lines = []
                
                def flush_list():
                    """Flush accumulated list items"""
                    nonlocal current_list_items, current_list_type
                    if current_list_items:
                        ol_type = '1' if current_list_type == 'numeric' else 'a'
                        list_html = f'<ol type="{ol_type}" style="margin-left: 20px; margin-top: 5px;">'
                        for item in current_list_items:
                            list_html += f'<li style="margin-bottom: 8px; line-height: 1.6;">{item}</li>'
                        list_html += '</ol>'
                        html_parts.append(list_html)
                        current_list_items = []
                        current_list_type = None
                
                # Process lines
                for ln in lines:
                    ln_stripped = ln.strip()
                    
                    # Check for empty line
                    if not ln_stripped:
                        # Empty line - end current list (if any)
                        if current_list_items:
                            flush_list()
                        prev_was_empty = True
                        continue
                    
                    # Try patterns: (1), 1., a., (a)
                    m_paren = re.match(r'^\((\d+|[a-z])\)\s+(.*)$', ln_stripped)  # (1) or (a)
                    m_dot = re.match(r'^(\d+|[a-z])\.\s+(.*)$', ln_stripped)     # 1. or a.
                    
                    # Check if this is a list introduction line
                    if list_intro_pat.match(ln_stripped):
                        # Flush current list (if any)
                        flush_list()
                        # Add this as paragraph (intro text)
                        flush_paragraph()
                        current_paragraph_lines.append(ln_stripped)
                        flush_paragraph()
                        prev_was_empty = False
                        continue
                    
                    if m_paren or m_dot:
                        # This is a list item - flush any pending paragraph first
                        flush_paragraph()
                        
                        # Extract item content and number
                        if m_paren:
                            num_str = m_paren.group(1)
                            content = m_paren.group(2).strip()
                        else:
                            num_str = m_dot.group(1)
                            content = m_dot.group(2).strip()
                        
                        # Detect if numeric or alpha
                        is_numeric = num_str.isdigit()
                        
                        # If number is "1" or "a" and we already have items, this might be a new list
                        if (num_str == '1' or num_str == 'a') and current_list_items:
                            # Check if type changed - if so, start new list
                            type_changed = (is_numeric and current_list_type != 'numeric') or (not is_numeric and current_list_type != 'alpha')
                            if type_changed:
                                flush_list()
                        
                        # Set or maintain list type
                        if current_list_type is None:
                            current_list_type = 'numeric' if is_numeric else 'alpha'
                        
                        # Add to current list
                        current_list_items.append(content)
                        prev_was_empty = False
                    else:
                        # Not a list item
                        # If previous line was empty OR no active list, this is a new paragraph
                        # Otherwise, it's continuation of last list item (same paragraph)
                        if prev_was_empty or not current_list_items:
                            # Start/continue paragraph
                            flush_list()  # Ensure any list is closed first
                            current_paragraph_lines.append(ln_stripped)
                        else:
                            # Continuation of last list item (no empty line between)
                            current_list_items[-1] += ' ' + ln_stripped
                        prev_was_empty = False
                
                # Flush remaining content
                flush_paragraph()
                flush_list()
                
                if not html_parts:
                    # fallback: convert newlines to <br>
                    html_parts.append('<p>' + '<br>'.join([l.strip() for l in lines if l.strip()]) + '</p>')
                
                return ''.join(html_parts)

            def flush_huruf():
                nonlocal pending_huruf_lines
                if pending_huruf_lines:
                    result_lines.extend(pending_huruf_lines)
                    pending_huruf_lines = None
            
            def flush_ayat():
                nonlocal pending_ayat_lines
                # Ensure huruf is flushed before ayat
                flush_huruf()
                if pending_ayat_lines:
                    result_lines.extend(pending_ayat_lines)
                    pending_ayat_lines = None
            
            def flush_pasal():
                nonlocal pending_pasal_lines
                # Ensure ayat (and huruf) are flushed before pasal
                flush_ayat()
                if pending_pasal_lines:
                    result_lines.extend(pending_pasal_lines)
                    pending_pasal_lines = None
            
            for line in lines:
                stripped = line.strip()
                
                # Stop processing when we hit Penjelasan section
                if penjelasan_section_pat.match(stripped):
                    # flush any pending markers before leaving main content
                    flush_pasal()
                    in_penjelasan_section = True
                    _logger.info("\n[OK] Reached Penjelasan section, stopping insertion")
                    break
                
                # PRIORITY 1: Insert Penjelasan Umum AFTER "DENGAN RAHMAT ... PRESIDEN REPUBLIK INDONESIA"
                # This should appear BEFORE "Menimbang"
                dengan_rahmat_match = dengan_rahmat_presiden_pat.match(stripped)
                if dengan_rahmat_match and not umum_inserted:
                    # First, add the "DENGAN RAHMAT ... PRESIDEN" line
                    result_lines.append(line)
                    # Then insert Penjelasan Umum right after
                    if 'umum' in penjelasan_map:
                        _logger.info(f"  -> TRIGGERING insertion after DENGAN RAHMAT...PRESIDEN")
                        _logger.info(f"  -> Penjelasan Umum content length: {len(penjelasan_map['umum'])} chars")
                        result_lines.extend([
                            '',
                            '📘 [Penjelasan Umum]:',
                            render_penjelasan_text(penjelasan_map['umum']),
                            ''
                        ])
                        umum_inserted = True
                        _logger.info(f"  -> Inserted Penjelasan Umum after DENGAN RAHMAT...PRESIDEN ({len(penjelasan_map['umum'])} chars)")
                    else:
                        _logger.warning(f"  -> Found DENGAN RAHMAT...PRESIDEN but 'umum' NOT in penjelasan_map!")
                    continue
                
                # PRIORITY 2: Detect "Menimbang" (fallback if PRESIDEN not found)
                menimbang_match = menimbang_pat.match(stripped)
                
                if menimbang_match and not umum_inserted:
                    # Insert Penjelasan Umum BEFORE "Menimbang"
                    if 'umum' in penjelasan_map:
                        _logger.info(f"  -> TRIGGERING insertion before Menimbang")
                        _logger.info(f"  -> Penjelasan Umum content length: {len(penjelasan_map['umum'])} chars")
                        result_lines.extend([
                            '',
                            '📘 [Penjelasan Umum]:',
                            render_penjelasan_text(penjelasan_map['umum']),
                            ''
                        ])
                        umum_inserted = True
                        _logger.info(f"  -> Inserted Penjelasan Umum before Menimbang ({len(penjelasan_map['umum'])} chars)")
                    else:
                        _logger.warning(f"  -> Found Menimbang but 'umum' NOT in penjelasan_map!")
                    # Now add the Menimbang line
                    result_lines.append(line)
                    continue
                
                # Check for BAB/BAGIAN headers - these should trigger flush of pending pasal
                if bab_header_pat.match(stripped):
                    flush_pasal()
                    result_lines.append(line)
                    continue
                
                # Track current Pasal in main content
                pasal_match = pasal_main_pat.match(stripped)
                if pasal_match:
                    # entering a new pasal: flush previous pending markers
                    if current_pasal is not None:
                        flush_pasal()
                    
                    current_pasal = pasal_match.group(1)
                    current_ayat = None
                    result_lines.append(line)
                    _logger.info(f"  -> Detected Pasal {current_pasal} in main content")
                    
                    # Prepare Pasal-level penjelasan to be inserted AFTER pasal content
                    key = f'pasal_{current_pasal}'
                    if key in penjelasan_map:
                        pending_pasal_lines = [
                            '',
                            f'💡 [Penjelasan Pasal {current_pasal}]:',
                            f'{render_penjelasan_text(penjelasan_map[key])}',
                            ''
                        ]
                        _logger.info(f"  -> Queued penjelasan for Pasal {current_pasal}")
                    else:
                        pending_pasal_lines = None
                    continue
                
                # Track current Ayat
                ayat_match = ayat_main_pat.match(stripped)
                if ayat_match and current_pasal:
                    # IMPORTANT: Before entering first Ayat, flush pending Pasal-level penjelasan
                    # This ensures Penjelasan Pasal appears AFTER Pasal header but BEFORE Ayat (1)
                    if current_ayat is None and pending_pasal_lines:
                        # This is the FIRST ayat in this Pasal - flush Pasal penjelasan now
                        result_lines.extend(pending_pasal_lines)
                        pending_pasal_lines = None
                        _logger.info(f"  -> Inserted Penjelasan Pasal {current_pasal} before first Ayat")
                    
                    # entering a new ayat: flush previous pending huruf/ayat
                    if current_ayat is not None:
                        flush_ayat()
                    
                    current_ayat = ayat_match.group(1)
                    result_lines.append(line)
                    
                    # Prepare Ayat-level penjelasan to be inserted AFTER ayat content
                    key = f'pasal_{current_pasal}_ayat_{current_ayat}'
                    if key in penjelasan_map:
                        pending_ayat_lines = [
                            '',
                            f'💡 [Penjelasan Ayat ({current_ayat})]:',
                            f'{render_penjelasan_text(penjelasan_map[key])}',
                            ''
                        ]
                        _logger.info(f"  -> Queued penjelasan for Pasal {current_pasal} Ayat ({current_ayat})")
                    else:
                        pending_ayat_lines = None
                    continue
                
                # Track Huruf (only if inside numbered list/ayat context)
                if current_ayat:
                    huruf_match = huruf_main_pat.match(stripped)
                    if huruf_match:
                        # entering a new huruf: flush previous pending huruf first
                        flush_huruf()
                        huruf = huruf_match.group(1)
                        result_lines.append(line)
                        
                        # Insert Huruf-level penjelasan IMMEDIATELY after huruf line
                        key = f'pasal_{current_pasal}_ayat_{current_ayat}_huruf_{huruf}'
                        if key in penjelasan_map:
                            result_lines.extend([
                                f'   💡 [Penjelasan Huruf {huruf}]:',
                                f'   {render_penjelasan_text(penjelasan_map[key])}',
                                ''
                            ])
                            _logger.info(f"  -> Inserted penjelasan immediately for Pasal {current_pasal} Ayat ({current_ayat}) Huruf {huruf}")
                        continue
                
                # Regular line - just add it
                result_lines.append(line)
            
            # Flush any remaining pending markers at end of main content
            flush_pasal()
            
            _logger.info(f"\n[OK] Penjelasan insertion complete")
            _logger.info("=" * 80)
            
            return '\n'.join(result_lines)
            
        except Exception as e:
            _logger.warning(f"Failed to insert penjelasan: {e}")
            return text

    def _extract_text_from_txt(self, txt_data):
        """Extract and format text from TXT upload (base64-encoded)."""
        try:
            # Decode base64 TXT data to bytes
            txt_bytes = base64.b64decode(txt_data)

            # DETECT: If the "TXT" file is actually a PDF binary
            if txt_bytes[:5] == b'%PDF-':
                _logger.warning("file_txt contains PDF binary data, extracting plain text via PyMuPDF...")
                try:
                    import fitz
                    doc = fitz.open(stream=txt_bytes, filetype="pdf")
                    text = "\n".join([page.get_text("text") for page in doc])
                    doc.close()
                    _logger.info(f"Extracted {len(text)} chars plain text from PDF binary in file_txt")
                except ImportError:
                    _logger.error("PyMuPDF (fitz) not installed! Cannot extract text from PDF binary in file_txt")
                    return '<p class="text-danger"><strong>Error:</strong> PyMuPDF tidak terinstall untuk mengekstrak PDF.</p>'
            else:
                text = None
                # Try common encodings
                for enc in ('utf-8', 'utf-16', 'cp1252', 'latin-1'):
                    try:
                        text = txt_bytes.decode(enc)
                        break
                    except Exception:
                        continue
                if text is None:
                    # Fallback with replacement
                    text = txt_bytes.decode('utf-8', errors='replace')

            # Normalize newlines
            text = text.replace('\r\n', '\n').replace('\r', '\n')

            # EARLY SNIPPET (raw before fixes) for Perubahan 7 block if present
            early_idx = text.lower().find('ketentuan ayat (2), ayat (3), ayat (5), ayat (6), ayat (7), dan ayat (8) pasal 43')
            if early_idx != -1:
                early_snippet = text[early_idx:early_idx+500]
                _logger.info('  [Perubahan7 RAW before OCR fixes] ' + early_snippet.replace('\n', ' | '))

            # Clean common OCR errors from PDF-to-TXT conversion
            text = re.sub(r'2g([A-Z])', r'28\1', text)  # Fix "2gD" -> "28D"
            text = re.sub(r'20l6', r'2016', text)  # Fix "20l6" -> "2016"
            text = re.sub(r'20l([0-9])', r'201\1', text)  # Fix "20l1" -> "2011"
            text = text.replace('Undang_Undang', 'Undang-Undang')
            text = text.replace('Undang_ Undang', 'Undang-Undang')
            
            # Fix Pasal number OCR errors FIRST (before any other processing)
            # This is CRITICAL: must be done early so other patterns can recognize "Pasal 45A", "Pasal 45B" etc.
            _logger.info("Fixing Pasal number OCR errors (458->45B, 45El->45B, etc.)...")
            text = re.sub(r'\bPasal\s+458\b', 'Pasal 45B', text, flags=re.IGNORECASE)  # 458 -> 45B
            text = re.sub(r'\bpasal\s+458\b', 'Pasal 45B', text, flags=re.IGNORECASE)  # pasal 458 -> Pasal 45B
            text = re.sub(r'\bPasal\s+45El\b', 'Pasal 45B', text, flags=re.IGNORECASE)  # 45El -> 45B
            text = re.sub(r'\bpasal\s+45El\b', 'Pasal 45B', text, flags=re.IGNORECASE)  # pasal 45El -> Pasal 45B
            text = re.sub(r'\bPasal\s+45B\b', 'Pasal 45B', text, flags=re.IGNORECASE)  # Normalize case
            text = re.sub(r'\bPasal\s+457\b', 'Pasal 45A', text, flags=re.IGNORECASE)  # 457 -> 45A (if exists)
            text = re.sub(r'\bpasal\s+457\b', 'Pasal 45A', text, flags=re.IGNORECASE)  # pasal 457 -> Pasal 45A
            text = re.sub(r'\bPasal\s+45A\b', 'Pasal 45A', text, flags=re.IGNORECASE)  # Normalize case
            
            # Fix common word OCR errors
            text = re.sub(r'\bayal\b', 'ayat', text, flags=re.IGNORECASE)  # Fix "ayal" -> "ayat"
            text = re.sub(r'\bs6lagaimana\b', 'sebagaimana', text, flags=re.IGNORECASE)  
            text = re.sub(r'\bElektronik dal\b', 'Elektronik dan', text)  
            text = re.sub(r'\bdal rekam\b', 'dan rekam', text, flags=re.IGNORECASE)  
            # Fix "dimaksud" OCR variants where 'maksud' middle chars misread (L/I/1)
            text = re.sub(r'\bdima[LI1]sud\b', 'dimaksud', text, flags=re.IGNORECASE)  # dimaLsud/dima1sud/dimaIsud -> dimaksud
            
            # Fix "cyber bullying" OCR errors - MUST BE BEFORE parenthesis fixes
            # "bullyingl" is actually "bullying)" where "l" is misread closing parenthesis
            text = re.sub(r'\(cgber\s+bullyingl', '(cyber bullying)', text, flags=re.IGNORECASE)  # (cgber bullyingl -> (cyber bullying)
            text = re.sub(r'cgber', 'cyber', text, flags=re.IGNORECASE)  # cgber -> cyber (fallback)
            text = re.sub(r'bullyingl', 'bullying)', text, flags=re.IGNORECASE)  # bullyingl -> bullying) (fallback)
            
            # Fix "huruf" followed by letter without space (e.g., "hurufj" -> "huruf j")
            text = re.sub(r'\bhuruf([a-z])\b', r'huruf \1', text, flags=re.IGNORECASE)  # hurufj -> huruf j
            _logger.info("Fixed merged 'huruf+letter' patterns (hurufj -> huruf j)")
            
            # Fix OCR error in privacy rights: "piuacg nqfus" -> "privacy rights"
            text = re.sub(r'\bpiuacg\s+nqfus\b', 'privacy rights', text, flags=re.IGNORECASE)
            
            # Fix parenthesis errors: "l" at end of numbers in parenthesis -> ")"
            text = re.sub(r'\((\d+)l\b', r'(\1)', text)  # (1l -> (1)
            text = re.sub(r'\((\d+[a-z])l\b', r'(\1)', text)  # (2al -> (2a)
            text = re.sub(r'\((\d+)1\b', r'(\1)', text)  # (11 -> (1)
            text = re.sub(r'\((\d+[a-z])1\b', r'(\1)', text)  # (2a1 -> (2a)
            
            # Fix double-l in parenthesis: (ll -> (1)
            text = re.sub(r'\(ll\b', r'(1', text)  # (ll -> (1
            text = re.sub(r'ayat\s+\(ll\)', 'ayat (1)', text, flags=re.IGNORECASE)  # ayat (ll) -> ayat (1)
            
            # Fix single lowercase L or uppercase I in parenthesis: (l) -> (1), (I) -> (1)
            text = re.sub(r'\([lI]\)', '(1)', text)  # (l) -> (1), (I) -> (1)
            text = re.sub(r'ayat\s+\([lI]\)', 'Ayat (1)', text, flags=re.IGNORECASE)  # Ayat (l) -> Ayat (1), Ayat (I) -> Ayat (1)
            text = re.sub(r'Pasal\s+(\d+)\s+ayat\s+\([lI]\)', r'Pasal \1 ayat (1)', text, flags=re.IGNORECASE)  # Pasal X ayat (l) -> Pasal X ayat (1)
            
            # Fix lowercase 'r' misread as '1' in parenthesis: (r) -> (1)
            text = re.sub(r'\(r\)', '(1)', text)  # (r) -> (1)
            text = re.sub(r'ayat\s+\(r\)', 'ayat (1)', text, flags=re.IGNORECASE)  # ayat (r) -> ayat (1)
            text = re.sub(r'Pasal\s+(\d+)\s+ayat\s+\(r\)', r'Pasal \1 ayat (1)', text, flags=re.IGNORECASE)  # Pasal X ayat (r) -> Pasal X ayat (1)
            
            # Fix incomplete parenthesis: ayat (1 -> ayat (1)
            text = re.sub(r'ayat\s+\((\d+[a-z]?)\s+', r'ayat (\1) ', text, flags=re.IGNORECASE)  # ayat (1 something -> ayat (1) something
            text = re.sub(r'ayat\s+\((\d+[a-z]?)\n', r'ayat (\1)\n', text, flags=re.IGNORECASE)  # ayat (1 newline -> ayat (1) newline

            # Fix broken enumeration line breaks: e.g.
            # "Ketentuan ayat (2), ayat (3), ayat (5), ayat (6), ayat (7), dan ayat\n(8)" -> single line
            _logger.info("Applying ayat enumeration line-break fixes...")
            # Join newline between 'ayat' (or 'dan ayat') and '(NUMBER)'
            text = re.sub(r'(ayat)\s*\n\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            text = re.sub(r'(dan\s+ayat)\s*\n\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            # Also handle optional preceding comma before newline
            text = re.sub(r'(ayat\s*\([^\)]+\),\s+dan\s+ayat)\s*\n\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            # Handle pattern: 'ayat (7), dan ayat\n(8) Pasal' -> 'ayat (7), dan ayat (8) Pasal'
            text = re.sub(r'(ayat\s*\(\d+[a-z]?\)\s*,?\s*dan\s+ayat)\s*\n\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            # Handle pattern: 'di antara ayat (7) dan ayat\n(8) Pasal' -> 'di antara ayat (7) dan ayat (8) Pasal'
            text = re.sub(r'(di\s+antara\s+ayat\s*\(\d+[a-z]?\)\s+dan\s+ayat)\s*\n\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            # Generic safeguard: line ends with 'dan ayat' and next line starts '(NUMBER) Pasal'
            text = re.sub(r'(dan\s+ayat)\s*\n\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            # Capitalize 'sehingga pasal 43' -> 'sehingga Pasal 43'
            text = re.sub(r'(?i)sehingga\s+pasal\s+(\d+)', r'sehingga Pasal \1', text)

            # Fix broken huruf enumeration: "c.\n" followed by text on next line -> "c. text" on same line
            # This handles cases where letter enumeration is separated from its content
            _logger.info("Fixing broken huruf (a., b., c., etc.) enumeration line breaks...")
            # Match pattern: line with just "c." or "c. " followed by newline and content
            before_fix = text
            text = re.sub(r'^([a-z])\.\s*$\n\s*(.+)', r'\1. \2', text, flags=re.MULTILINE)
            if text != before_fix:
                _logger.info("  [OK] Fixed broken huruf enumeration (standalone letter on one line, content on next)")
            
            # FIX: Separate "Pasal 1 Dalam Undang-Undang..." into two lines
            # Also handle Roman numerals like "Pasal I Dalam Undang-Undang..."
            # "Pasal 1 Dalam Undang-Undang ini yang dimaksud dengan:" -> "Pasal 1\nDalam Undang-Undang ini yang dimaksud dengan:"
            _logger.info("Separating 'Dalam Undang-Undang' subtitle from Pasal header...")
            
            # Check if pattern exists before replacement (support both numeric and Roman)
            pasal_dalam_pattern = re.compile(r'\b(Pasal\s+(?:\d+[A-Z]?|[IVX]+))\s+(Dalam\s+Undang-Undang\s+ini\s+yang\s+dimaksud\s+dengan:)', re.IGNORECASE)
            matches = pasal_dalam_pattern.findall(text)
            if matches:
                _logger.info(f"  Found {len(matches)} occurrences of 'Pasal X Dalam Undang-Undang...'")
                for match in matches[:3]:  # Show first 3
                    _logger.info(f"    - '{match[0]} {match[1][:50]}...'")
            else:
                _logger.info("  No 'Pasal X Dalam Undang-Undang...' pattern found")
            
            # Separate for numeric Pasal (Pasal 1, Pasal 2, etc.)
            text = re.sub(
                r'\b(Pasal\s+\d+[A-Z]?)\s+(Dalam\s+Undang-Undang\s+ini\s+yang\s+dimaksud\s+dengan:)',
                r'\1\n\2',
                text,
                flags=re.IGNORECASE
            )
            
            # Separate for Roman numeral Pasal (Pasal I, Pasal II, etc.)
            text = re.sub(
                r'\b(Pasal\s+[IVX]+)\s+(Dalam\s+Undang-Undang\s+ini\s+yang\s+dimaksud\s+dengan:)',
                r'\1\n\2',
                text,
                flags=re.IGNORECASE
            )
            
            # Log snippet to verify separation worked
            if matches:
                # Find "Pasal 1" or "Pasal I" in text after separation
                pasal1_idx = text.lower().find('pasal 1')
                if pasal1_idx == -1:
                    pasal1_idx = text.lower().find('pasal i')
                if pasal1_idx != -1:
                    snippet = text[pasal1_idx:pasal1_idx+150]
                    _logger.info(f"  After separation, Pasal area: {snippet.replace(chr(10), ' | ')}")

            
            # CRITICAL: Separate "Pasal 45A", "Pasal 45B", etc. onto their own lines
            # But ONLY when they appear as actual section headers (after punctuation + followed by content)
            # NOT when mentioned in reference lists like "pasal 25A, pasal 28D"
            _logger.info("Separating Pasal headers with letter suffixes (45A, 45B, etc.) from following content...")
            
            # BROAD APPROACH: Separate "Pasal 45B" when it appears after any sentence ending
            # Pattern matches: ). or .) or closing paren followed by period
            
            # Match various patterns of sentence ending before Pasal 45A/45B
            # Examples: "rupiah). Pasal 45B", "000,00 (satu miliar rupiah). Pasal 45B", "Pasal 29 dipidana. Pasal 45B"
            
            text = re.sub(
                r'([.)])(\s*)Pasal\s+45B\s+',
                r'\1\nPasal 45B\n',
                text,
                flags=re.IGNORECASE
            )
            
            text = re.sub(
                r'([.)])(\s*)Pasal\s+45A\s+',
                r'\1\nPasal 45A\n',
                text,
                flags=re.IGNORECASE
            )
            
            # Also handle if there's a newline already between ). and Pasal
            text = re.sub(
                r'([.)])\s*\n\s*Pasal\s+45B\s+',
                r'\1\nPasal 45B\n',
                text,
                flags=re.IGNORECASE
            )
            
            # Merge broken lines from PDF conversion
            _logger.info("Starting sentence merge process...")
            text = self._merge_broken_lines(text)

            # SNIPPET after merge
            mid_idx = text.lower().find('ketentuan ayat (2), ayat (3), ayat (5), ayat (6), ayat (7), dan ayat (8) pasal 43')
            if mid_idx != -1:
                mid_snippet = text[mid_idx:mid_idx+550]
                _logger.info('  [Perubahan7 AFTER merge] ' + mid_snippet.replace('\n', ' | '))

            # Guarantee continuity for Perubahan 7 full sentence (avoid truncation after "; serta")
            # Target expected tail phrase
            expected_tail_pat = re.compile(r'serta\s+penjelasan\s+ayat\s*\(1\)\s+Pasal\s+43\s+diubah\s+sehingga\s+Pasal\s+43\s+berbunyi\s+sebagai\s+berikut:', re.IGNORECASE)
            # If we have the beginning but not the tail, try to locate tail elsewhere and append
            start_pat = re.compile(r'Ketentuan\s+ayat\s*\(2\).*?yakni\s+ayat\s*\(7a\);\s*serta\b', re.IGNORECASE | re.DOTALL)
            has_start = bool(start_pat.search(text))
            has_tail = bool(expected_tail_pat.search(text))
            if has_start and has_tail:
                # Ensure they are contiguous (no accidental removal). If a newline splits, remove it.
                text = re.sub(r'(yakni\s+ayat\s*\(7a\);)\s*\n\s*(serta\s+penjelasan)', r'\1 \2', text, flags=re.IGNORECASE)
            elif has_start and not has_tail:
                _logger.warning('  Perubahan7: tail phrase missing after "; serta" — attempting recovery.')
                # Attempt heuristic recovery: look for a truncated 'serta' ending
                # If next 200 chars after the start pattern end with 'serta', log for diagnostics
                m = start_pat.search(text)
                if m:
                    after_pos = m.end()
                    diagnostic_tail = text[after_pos:after_pos+250]
                    _logger.warning('  [Perubahan7 DIAGNOSTIC after start] ' + diagnostic_tail.replace('\n', ' | '))
            else:
                _logger.info('  Perubahan7: start pattern not detected or already intact.')

            # SECOND PASS: enumeration newline fixes after line-merging (merge may reintroduce breaks)
            _logger.info("Re-applying ayat enumeration fixes post-merge (spasi-aware)...")
            pre_count_generic = len(re.findall(r'ayat\s*\n\s*\(\s*\d', text, flags=re.IGNORECASE))
            pre_count_dan = len(re.findall(r'dan\s+ayat\s*\n\s*\(\s*\d', text, flags=re.IGNORECASE))
            if pre_count_generic or pre_count_dan:
                _logger.info(f"  Found {pre_count_generic} generic and {pre_count_dan} 'dan ayat' broken enumerations before fix")
            # Generic ayat followed by newline + (number)
            text = re.sub(r'(ayat)\s*\n\s*\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            # 'dan ayat' case
            text = re.sub(r'(dan\s+ayat)\s*\n\s*\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            # Enumeration with preceding list and comma
            text = re.sub(r'(ayat\s*\([^\)]+\),\s+dan\s+ayat)\s*\n\s*\(\s*(\d+[a-z]?)\)', r'\1 (\2)', text, flags=re.IGNORECASE)
            # Pattern ending before 'Pasal'
            text = re.sub(r'(ayat\s*\(\d+[a-z]?\)\s*,?\s*dan\s+ayat)\s*\n\s*\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            # 'di antara ayat (7) dan ayat\n(8) Pasal'
            text = re.sub(r'(di\s+antara\s+ayat\s*\(\d+[a-z]?\)\s+dan\s+ayat)\s*\n\s*\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            # Simple trailing 'dan ayat' before Pasal
            text = re.sub(r'(dan\s+ayat)\s*\n\s*\(\s*(\d+[a-z]?)\s*(Pasal)\b', r'\1 (\2) \3', text, flags=re.IGNORECASE)
            post_count_generic = len(re.findall(r'ayat\s*\n\s*\(\s*\d', text, flags=re.IGNORECASE))
            post_count_dan = len(re.findall(r'dan\s+ayat\s*\n\s*\(\s*\d', text, flags=re.IGNORECASE))
            _logger.info(f"  Remaining broken enumerations after fix: generic={post_count_generic}, dan={post_count_dan}")

            # Log snippet around Perubahan 7 if present for debugging
            perubahan_idx = text.lower().find('ketentuan ayat (2), ayat (3), ayat (5), ayat (6), ayat (7), dan ayat')
            if perubahan_idx != -1:
                snippet = text[perubahan_idx:perubahan_idx+320]
                _logger.info("  Snippet Perubahan 7 after fixes: " + snippet.replace('\n', ' | '))

            # Ensure continuity: keep 'yakni ayat (7a); serta penjelasan ...' on ONE line
            # Remove unwanted newline if previously inserted
            continuity_before = len(re.findall(r'yakni\s+ayat\s*\(7a\);\s*\n\s*serta\s+penjelasan', text, flags=re.IGNORECASE))
            # First: join 'yakni ayat (7a);' with 'serta'
            text = re.sub(r'(yakni\s+ayat\s*\(7a\);)\s*\n\s*(serta)', r'\1 \2', text, flags=re.IGNORECASE)
            # Second: join 'serta' with 'penjelasan' (in case separated)
            text = re.sub(r'(;\s*serta)\s*\n\s*(penjelasan\s+ayat)', r'\1 \2', text, flags=re.IGNORECASE)
            # Third: full pattern match
            text = re.sub(r'(yakni\s+ayat\s*\(7a\);)\s*\n\s*(serta\s+penjelasan\s+ayat\s*\(1\)\s+Pasal\s+43)', r'\1 \2', text, flags=re.IGNORECASE)
            # Also normalize multiple spaces
            text = re.sub(r'(yakni\s+ayat\s*\(7a\);)\s{2,}(serta\s+penjelasan)', r'\1 \2', text, flags=re.IGNORECASE)
            continuity_after = len(re.findall(r'yakni\s+ayat\s*\(7a\);\s+serta\s+penjelasan', text, flags=re.IGNORECASE))
            if continuity_before or continuity_after:
                _logger.info(f"  Fixed 7a continuity: before_breaks={continuity_before}, after_continuous={continuity_after}")

            # Insert penjelasan into text if Penjelasan section exists
            # Check for both "PENJELASAN ATAS" and "Penjelasan" (case-insensitive)
            if 'PENJELASAN' in text.upper():
                try:
                    _logger.info("Detected Penjelasan section, starting insertion...")
                    text = self._insert_penjelasan_into_text(text)
                    _logger.info("Successfully inserted penjelasan into text")
                except Exception as e:
                    _logger.error(f"Failed to insert penjelasan: {e}", exc_info=True)
                    _logger.warning(f"Using original text without penjelasan integration")
            else:
                _logger.info("No Penjelasan section found in text")
            
            # Format to HTML using existing formatter
            formatted = self._format_text_to_html(text)
            return f'<div class="txt-content" style="line-height: 1.8;">{formatted}</div>'
        except Exception as e:
            _logger.error(f"Error extracting TXT text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari TXT: {str(e)}</p>'

    @api.onchange('file_docx')
    def _onchange_file_docx(self):
        """Auto-extract text from DOCX and convert to PDF"""
        if self.file_docx:
            _logger.info(f"DOCX uploaded for regulation {self.id or 'new'}")
            
            # 1. Extract text from DOCX (more accurate)
            extracted_text = self._extract_text_from_docx(self.file_docx)
            if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                self.isi_peraturan = extracted_text
                _logger.info(f"Successfully extracted text from DOCX ({len(extracted_text)} characters)")
            
            # 2. Convert DOCX to PDF for download
            pdf_data = self._convert_docx_to_pdf(self.file_docx)
            if pdf_data:
                self.file_pdf = pdf_data
                _logger.info("Successfully converted DOCX to PDF for download")
        else:
            # File DOCX dihapus - clear content dan PDF
            _logger.info(f"DOCX removed for regulation {self.id or 'new'}, clearing content")
            self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
            # Only clear PDF if it was auto-generated from DOCX
            if not self.file_pdf or self.file_pdf:
                # Check if we should clear PDF (don't clear if manually uploaded)
                pass  # Let user manage PDF separately

    @api.onchange('file_txt')
    def _onchange_file_txt(self):
        """Store TXT file without auto-overwriting existing content.
        User can manually re-extract using 'Re-Extract' button."""
        if self.file_txt:
            _logger.info(f"TXT file uploaded for regulation {self.id or 'new'} (stored without auto-extract)")
            # Only auto-extract if content is empty or default placeholder
            current = (self.isi_peraturan or '').strip()
            is_empty = not current or current == '<p>Isi peraturan belum tersedia</p>' or len(current) < 50
            if is_empty:
                extracted_text = self._extract_text_from_txt(self.file_txt)
                if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                    self.isi_peraturan = extracted_text
                    _logger.info(f"Auto-extracted text from TXT for empty record ({len(extracted_text)} characters)")
            else:
                _logger.info(f"Existing content preserved ({len(current)} chars). Use 'Re-Extract' button to overwrite.")
        else:
            # TXT removed - clear only if no PDF and no DOCX exists
            if not self.file_pdf and not self.file_docx:
                _logger.info(f"TXT removed and no other files present for regulation {self.id or 'new'}, clearing content")
                self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
    
    @api.onchange('file_pdf')
    def _onchange_file_pdf(self):
        """Store PDF file without auto-overwriting existing content.
        User can manually re-extract using 'Re-Extract' button."""
        if self.file_pdf:
            _logger.info(f"PDF file uploaded for regulation {self.id or 'new'} (stored without auto-extract)")
            # Only auto-extract if content is empty or default placeholder
            current = (self.isi_peraturan or '').strip()
            is_empty = not current or current == '<p>Isi peraturan belum tersedia</p>' or len(current) < 50
            if is_empty:
                extracted_text = self._extract_text_from_pdf(self.file_pdf)
                if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                    self.isi_peraturan = extracted_text
                    _logger.info(f"Auto-extracted text from PDF for empty record ({len(extracted_text)} characters)")
            else:
                _logger.info(f"Existing content preserved ({len(current)} chars). Use 'Re-Extract' button to overwrite.")
        else:
            # File PDF dihapus - clear content jika tidak ada DOCX dan TXT
            if not self.file_docx and not self.file_txt:
                _logger.info(f"PDF removed and no DOCX/TXT present for regulation {self.id or 'new'}, clearing content")
                self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
    
    @api.onchange('bentuk')
    def _onchange_bentuk(self):
        """Auto-generate bentuk_singkat dari bentuk"""
        if self.bentuk:
            # Mapping umum bentuk singkat
            bentuk_mapping = {
                'Undang-Undang': 'UU',
                'Peraturan Pemerintah': 'PP',
                'Peraturan Presiden': 'Perpres',
                'Peraturan Menteri': 'Permen',
                'Peraturan Daerah': 'Perda',
                'Keputusan Presiden': 'Keppres',
                'Keputusan Menteri': 'Kepmen',
                'Instruksi Presiden': 'Inpres',
                'Surat Edaran': 'SE',
            }
            
            for key, value in bentuk_mapping.items():
                if key.lower() in self.bentuk.lower():
                    self.bentuk_singkat = value
                    break
    
    @api.onchange('tanggal_penetapan')
    def _onchange_tanggal_penetapan(self):
        """Auto set tahun dari tanggal penetapan"""
        if self.tanggal_penetapan:
            self.tahun = self.tanggal_penetapan.year
            # Auto set tanggal berlaku jika belum diisi
            if not self.tanggal_berlaku:
                self.tanggal_berlaku = self.tanggal_penetapan
    
    # Computed Fields
    @api.depends('tipe_dokumen')
    def _compute_hierarchy_order(self):
        """Compute hierarchy order based on Indonesian legal hierarchy"""
        hierarchy_mapping = {
            'uud_1945': 1,           
            'tap_mpr': 2,            
            'undang_undang': 3,      
            'perpu': 4,              
            'peraturan_pemerintah': 5,  
            'peraturan_presiden': 6,    
            'keputusan_presiden': 7,    
            'instruksi_presiden': 8,    
            'peraturan_menteri': 9,     
            'keputusan_menteri': 10,    
            'peraturan_daerah': 11,     
            'peraturan_gubernur': 12,   
        }
        
        for record in self:
            record.hierarchy_order = hierarchy_mapping.get(record.tipe_dokumen, 999)

    @api.model
    def _auto_init(self):
        """Auto-init untuk memastikan kompatibilitas field baru"""
        # Call parent init first
        res = super()._auto_init()
        
        # Check and add missing fields for compatibility
        try:
            # List field baru yang mungkin belum ada
            new_fields = [
                ('isi_peraturan', 'TEXT'),
                ('kata_kunci', 'TEXT'), 
                ('ringkasan', 'TEXT'),
                ('hierarchy_order', 'INTEGER')
            ]
            
            for field_name, field_type in new_fields:
                # Check if field exists
                self._cr.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name=%s AND column_name=%s
                """, (self._table, field_name))
                
                if not self._cr.fetchone():
                    # Add missing field
                    if field_type == 'TEXT':
                        self._cr.execute(f"ALTER TABLE {self._table} ADD COLUMN {field_name} TEXT")
                    elif field_type == 'INTEGER':
                        self._cr.execute(f"ALTER TABLE {self._table} ADD COLUMN {field_name} INTEGER DEFAULT 999")
                    
                    print(f"[COMPATIBILITY] Added missing field: {field_name}")
            
            # Set default values for existing records with NULL values
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET hierarchy_order = 999 
                WHERE hierarchy_order IS NULL
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET isi_peraturan = '<p>Isi peraturan belum tersedia</p>' 
                WHERE isi_peraturan IS NULL OR isi_peraturan = ''
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET kata_kunci = '' 
                WHERE kata_kunci IS NULL
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET ringkasan = 'Ringkasan belum tersedia' 
                WHERE ringkasan IS NULL OR ringkasan = ''
            """)
            
        except Exception as e:
            # Jika ada error, log tapi jangan stop install
            print(f"[COMPATIBILITY WARNING] {str(e)}")
            
        return res

    @api.model
    def create(self, vals_list):
        """Override create untuk validasi dan auto-extract TXT/DOCX/PDF"""
        # Handle both single dict and list of dicts
        if isinstance(vals_list, dict):
            vals_list = [vals_list]
        
        for vals in vals_list:
            # Auto generate bentuk_singkat jika tidak ada
            if vals.get('bentuk') and not vals.get('bentuk_singkat'):
                bentuk = vals['bentuk']
                if 'Menteri' in bentuk:
                    vals['bentuk_singkat'] = 'Permen'
                elif 'Pemerintah' in bentuk:
                    vals['bentuk_singkat'] = 'PP'
                elif 'Presiden' in bentuk:
                    vals['bentuk_singkat'] = 'Perpres'
            
            # Priority: TXT > DOCX > PDF
            if vals.get('file_txt'):
                try:
                    if not vals.get('isi_peraturan'):
                        extracted_text = self._extract_text_from_txt(vals['file_txt'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info("Auto-extracted TXT text during creation")
                except Exception as e:
                    _logger.error(f"Failed to process TXT during creation: {e}")
            elif vals.get('file_docx'):
                try:
                    # Extract text from DOCX
                    if not vals.get('isi_peraturan'):
                        extracted_text = self._extract_text_from_docx(vals['file_docx'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info("Auto-extracted DOCX text during creation")
                    
                    # Convert DOCX to PDF for download
                    pdf_data = self._convert_docx_to_pdf(vals['file_docx'])
                    if pdf_data:
                        vals['file_pdf'] = pdf_data
                        _logger.info("Auto-converted DOCX to PDF during creation")
                except Exception as e:
                    _logger.error(f"Failed to process DOCX during creation: {e}")
            
            # Fallback: Extract from PDF if uploaded manually
            elif vals.get('file_pdf') and not vals.get('isi_peraturan'):
                try:
                    extracted_text = self._extract_text_from_pdf(vals['file_pdf'])
                    if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                        vals['isi_peraturan'] = extracted_text
                        _logger.info("Auto-extracted PDF text during creation")
                except Exception as e:
                    _logger.error(f"Failed to auto-extract PDF during creation: {e}")
        
        return super(LegalRegulation, self).create(vals_list)
    
    def write(self, vals):
        """Override write untuk auto-extract TXT/DOCX/PDF saat update"""
        # Check if files are being deleted (set to False)
        file_txt_deleted = 'file_txt' in vals and not vals.get('file_txt')
        file_docx_deleted = 'file_docx' in vals and not vals.get('file_docx')
        file_pdf_deleted = 'file_pdf' in vals and not vals.get('file_pdf')
        
        # Handle file deletion - clear content if all files removed
        if (file_txt_deleted or not self.file_txt) and (file_docx_deleted or not self.file_docx) and file_pdf_deleted:
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"All files deleted for regulation ID {self.id}, clearing content")
        elif file_txt_deleted and not vals.get('file_pdf') and not self.file_pdf and not self.file_docx:
            # TXT deleted and no other file exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"TXT deleted and no other files present for regulation ID {self.id}, clearing content")
        elif file_docx_deleted and not vals.get('file_pdf') and not self.file_pdf and not self.file_txt:
            # DOCX deleted and no PDF/TXT exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"DOCX deleted and no PDF/TXT present for regulation ID {self.id}, clearing content")
        elif file_pdf_deleted and not vals.get('file_docx') and not self.file_docx and not self.file_txt:
            # PDF deleted and no DOCX/TXT exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"PDF deleted and no DOCX/TXT present for regulation ID {self.id}, clearing content")
        
        # Priority: TXT > DOCX > PDF
        # Only auto-extract if isi_peraturan is NOT already set with real content
        if vals.get('file_txt') and not file_txt_deleted:
            try:
                if 'isi_peraturan' not in vals:
                    current_content = (self.isi_peraturan or '').strip()
                    is_empty = not current_content or current_content == '<p>Isi peraturan belum tersedia</p>' or len(current_content) < 50
                    if is_empty:
                        extracted_text = self._extract_text_from_txt(vals['file_txt'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info(f"Auto-extracted TXT text for regulation ID {self.id} (was empty)")
                    else:
                        _logger.info(f"File TXT uploaded for regulation ID {self.id} but existing content preserved ({len(current_content)} chars). Use 'Re-Extract' button.")
            except Exception as e:
                _logger.error(f"Failed to process TXT during write: {e}")
        elif vals.get('file_docx') and not file_docx_deleted:
            try:
                # Extract text from DOCX only if content is empty
                if 'isi_peraturan' not in vals:
                    current_content = (self.isi_peraturan or '').strip()
                    is_empty = not current_content or current_content == '<p>Isi peraturan belum tersedia</p>' or len(current_content) < 50
                    if is_empty:
                        extracted_text = self._extract_text_from_docx(vals['file_docx'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info(f"Auto-extracted DOCX text for regulation ID {self.id} (was empty)")
                    else:
                        _logger.info(f"File DOCX uploaded for regulation ID {self.id} but existing content preserved. Use 'Re-Extract' button.")
                
                # Convert DOCX to PDF for download
                pdf_data = self._convert_docx_to_pdf(vals['file_docx'])
                if pdf_data:
                    vals['file_pdf'] = pdf_data
                    _logger.info(f"Auto-converted DOCX to PDF for regulation ID {self.id}")
            except Exception as e:
                _logger.error(f"Failed to process DOCX during write: {e}")
        
        # Fallback: Extract from PDF if uploaded manually and content is empty
        elif vals.get('file_pdf') and not file_pdf_deleted and 'isi_peraturan' not in vals:
            try:
                current_content = (self.isi_peraturan or '').strip()
                is_empty = not current_content or current_content == '<p>Isi peraturan belum tersedia</p>' or len(current_content) < 50
                if is_empty:
                    extracted_text = self._extract_text_from_pdf(vals['file_pdf'])
                    if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                        vals['isi_peraturan'] = extracted_text
                        _logger.info(f"Auto-extracted PDF text for regulation ID {self.id} (was empty)")
                else:
                    _logger.info(f"File PDF uploaded for regulation ID {self.id} but existing content preserved. Use 'Re-Extract' button.")
            except Exception as e:
                _logger.error(f"Failed to auto-extract PDF during write: {e}")
        
        return super(LegalRegulation, self).write(vals)
    
    def name_get(self):
        """Custom name display"""
        result = []
        for record in self:
            if record.bentuk_singkat and record.nomor and record.tahun:
                name = f"{record.bentuk_singkat} No. {record.nomor}/{record.tahun}"
            else:
                name = record.judul[:50] + '...' if len(record.judul) > 50 else record.judul
            result.append((record.id, name))
        return result
    
    @api.model
    def _name_search(self, name, args=None, operator='ilike', limit=100, name_get_uid=None):
        """Enhanced search functionality"""
        args = args or []
        domain = []
        if name:
            domain = ['|', '|', '|', '|', '|', '|', '|', '|',
                     ('judul', operator, name),
                     ('nomor', operator, name),
                     ('bentuk_singkat', operator, name),
                     ('subjek', operator, name),
                     ('keterangan', operator, name),
                     ('isi_peraturan', operator, name),
                     ('kata_kunci', operator, name),
                     ('ringkasan', operator, name),
                     ('bentuk', operator, name)]
        
        regulation_ids = self._search(domain + args, limit=limit, access_rights_uid=name_get_uid)
        return self.browse(regulation_ids).name_get()
    
    def action_download_pdf(self):
        """Action untuk download file PDF"""
        self.ensure_one()
        if not self.file_pdf:
            from odoo.exceptions import UserError
            raise UserError('File PDF tidak tersedia untuk peraturan ini.')
        
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content?model=legal.regulation&id={self.id}&field=file_pdf&filename_field=file_name&download=true',
            'target': 'self',
        }
    
    def action_create_consolidation(self):
        """Buat konsolidasi V2 - langsung dari file TXT"""
        self.ensure_one()
        
        # Validasi
        if not self.file_txt:
            from odoo.exceptions import UserError
            raise UserError('Peraturan ini tidak memiliki file TXT. Upload file TXT terlebih dahulu di tab "File Dokumen".')
        
        if not self.perubahan_dari_ids:
            from odoo.exceptions import UserError
            raise UserError('Tidak ada perubahan yang tercatat. Tambahkan perubahan UU di tab "Perubahan UU" terlebih dahulu.')
        
        # Cek perubahan punya file TXT
        perubahan_valid = self.perubahan_dari_ids.filtered(
            lambda p: p.file_txt_perubahan or (p.peraturan_pengubah_id and p.peraturan_pengubah_id.file_txt)
        )
        
        if not perubahan_valid:
            from odoo.exceptions import UserError
            raise UserError('Tidak ada perubahan yang memiliki file TXT. Upload file TXT perubahan terlebih dahulu.')
        
        # Buat nama default
        default_name = f"Konsolidasi: {self.nama_lengkap}"
        
        # Buat record konsolidasi v2 langsung
        consolidation = self.env['legal.regulation.consolidation.v2'].create({
            'name': default_name,
            'peraturan_induk_id': self.id,
            'perubahan_ids': [(6, 0, perubahan_valid.ids)],
            'display_mode': 'annotated',
        })
        
        # Return action untuk buka form konsolidasi v2
        return {
            'name': 'Konsolidasi Peraturan',
            'type': 'ir.actions.act_window',
            'res_model': 'legal.regulation.consolidation.v2',
            'res_id': consolidation.id,
            'view_mode': 'form',
            'target': 'current',
        }
    
    def action_reextract_pdf(self):
        """Re-extract content from uploaded file (TXT/DOCX/PDF) with latest extraction method"""
        self.ensure_one()
        
        if not (self.file_txt or self.file_docx or self.file_pdf):
            from odoo.exceptions import UserError
            raise UserError('Tidak ada file untuk di-ekstrak ulang.')
        
        try:
            # Determine which file to extract from (priority: TXT > DOCX > PDF)
            if self.file_txt:
                _logger.info(f"Re-extracting TXT for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_txt(self.file_txt)
                file_type = "TXT"
            elif self.file_docx:
                _logger.info(f"Re-extracting DOCX for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_docx(self.file_docx)
                file_type = "DOCX"
            else:
                _logger.info(f"Re-extracting PDF for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_pdf(self.file_pdf)
                file_type = "PDF"
            
            # Embed extraction timestamp as HTML comment for cache-busting/debugging
            from odoo.tools import format_datetime
            ts = format_datetime(self.env, fields.Datetime.now())
            extracted_html = f"<!-- reextracted from {file_type}: {ts} -->\n" + (extracted_html or '')
            
            # Update the isi_peraturan field
            self.write({
                'isi_peraturan': extracted_html
            })
            
            _logger.info(f"Successfully re-extracted {file_type} for regulation ID {self.id}")
            
            # Reload the view to ensure the latest HTML is shown
            return {'type': 'ir.actions.client', 'tag': 'reload'}
            
        except Exception as e:
            _logger.error(f"Failed to re-extract file for regulation ID {self.id}: {e}")
            from odoo.exceptions import UserError
            raise UserError(f'Gagal mengekstrak ulang file: {str(e)}')

    @api.model
    def _auto_init(self):
        """Ensure compatibility during install - add missing fields safely"""
        res = super()._auto_init()
        
        try:
            # Check dan tambahkan field yang mungkin belum ada
            fields_to_check = [
                ('isi_peraturan', 'TEXT', '<p>Isi peraturan belum tersedia</p>'),
                ('kata_kunci', 'VARCHAR', ''),
                ('ringkasan', 'TEXT', 'Ringkasan belum tersedia'),
                ('hierarchy_order', 'INTEGER', '999')
            ]
            
            for field_name, field_type, default_value in fields_to_check:
                self._cr.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name='legal_regulation' AND column_name=%s
                """, (field_name,))
                
                if not self._cr.fetchone():
                    if field_type == 'TEXT':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} TEXT")
                        if default_value:
                            self._cr.execute(f"UPDATE legal_regulation SET {field_name} = %s WHERE {field_name} IS NULL", (default_value,))
                    elif field_type == 'VARCHAR':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} VARCHAR")
                        if default_value:
                            self._cr.execute(f"UPDATE legal_regulation SET {field_name} = %s WHERE {field_name} IS NULL", (default_value,))
                    elif field_type == 'INTEGER':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} INTEGER DEFAULT {default_value}")
                        
            self._cr.commit()
            
        except Exception as e:
            # Jika ada error, log tapi jangan crash
            import logging
            _logger = logging.getLogger(__name__)
            _logger.warning(f"Field migration error: {e}")
            
        return res


class LegalRegulationType(models.Model):
    _name = 'legal.regulation.type'
    _description = 'Legal Regulation Type'
    _order = 'sequence, name'
    
    name = fields.Char('Nama Tipe', required=True)
    code = fields.Char('Kode', required=True)
    description = fields.Text('Deskripsi')
    sequence = fields.Integer('Urutan', default=10)
    active = fields.Boolean('Active', default=True)
    
    _sql_constraints = [
        ('code_unique', 'unique(code)', 'Kode tipe peraturan harus unik!')
    ]